"""
Document Processor
Version: 1.0.0
Date: 2026-01-05

Purpose: Process tender documents (.docx, .pdf, .txt) for mention-based proposals.
"""

from pathlib import Path
from typing import Optional, List, Dict, Any
import tempfile
import shutil

from .utils import get_logger

logger = get_logger(__name__)


class DocumentProcessor:
    """Process tender documents and extract text."""

    @staticmethod
    def process_document(file_path: Path) -> str:
        """
        Process document and extract text.

        Args:
            file_path: Path to document

        Returns:
            Extracted text

        Raises:
            ValueError: If unsupported file type
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix == '.txt':
            return DocumentProcessor._process_txt(file_path)
        elif suffix == '.docx':
            return DocumentProcessor._process_docx(file_path)
        elif suffix == '.pdf':
            return DocumentProcessor._process_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    @staticmethod
    def _process_txt(file_path: Path) -> str:
        """Process .txt file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            logger.info(f"Processed TXT: {file_path.name} ({len(text)} chars)")
            return text

        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()

            logger.info(f"Processed TXT (latin-1): {file_path.name} ({len(text)} chars)")
            return text

    @staticmethod
    def _process_docx(file_path: Path) -> str:
        """Process .docx file."""
        try:
            from docx import Document

            doc = Document(file_path)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)

            text = '\n\n'.join(paragraphs)

            logger.info(f"Processed DOCX: {file_path.name} ({len(text)} chars, {len(doc.paragraphs)} paragraphs)")
            return text

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Failed to process DOCX: {e}")
            raise

    @staticmethod
    def _process_pdf(file_path: Path) -> str:
        """Process .pdf file."""
        try:
            import PyPDF2

            text_parts = []

            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)

                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)

            text = '\n\n'.join(text_parts)

            logger.info(f"Processed PDF: {file_path.name} ({len(text)} chars, {len(pdf_reader.pages)} pages)")
            return text

        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            raise
        except Exception as e:
            logger.error(f"Failed to process PDF: {e}")
            raise

    @staticmethod
    def save_document_with_mentions(
        text: str,
        output_path: Path,
        output_format: str = 'txt'
    ) -> bool:
        """
        Save document with @mentions.

        Args:
            text: Document text with @mentions
            output_path: Output file path
            output_format: Format: txt, docx

        Returns:
            True if successful
        """
        output_path = Path(output_path)

        if output_format == 'txt':
            return DocumentProcessor._save_txt(text, output_path)
        elif output_format == 'docx':
            return DocumentProcessor._save_docx(text, output_path)
        else:
            raise ValueError(f"Unsupported output format: {output_format}")

    @staticmethod
    def _save_txt(text: str, output_path: Path) -> bool:
        """Save as .txt file."""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)

            logger.info(f"Saved TXT: {output_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to save TXT: {e}")
            return False

    @staticmethod
    def _save_docx(text: str, output_path: Path) -> bool:
        """Save as .docx file."""
        try:
            from docx import Document

            doc = Document()

            # Split text into paragraphs and add to document
            for para_text in text.split('\n\n'):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())

            doc.save(output_path)

            logger.info(f"Saved DOCX: {output_path}")
            return True

        except ImportError:
            logger.error("python-docx not installed")
            return False
        except Exception as e:
            logger.error(f"Failed to save DOCX: {e}")
            return False

    @staticmethod
    def extract_sections(text: str) -> Dict[str, str]:
        """
        Extract sections from document based on headings.

        Args:
            text: Document text

        Returns:
            Dict of section_name -> section_text

        Example:
            >>> text = "SECTION 1: INTRO\\nContent here\\n\\nSECTION 2: DETAILS\\nMore content"
            >>> sections = DocumentProcessor.extract_sections(text)
            >>> sections['SECTION 1']
            'Content here'
        """
        sections = {}
        current_section = "HEADER"
        current_content = []

        lines = text.split('\n')

        for line in lines:
            # Check if this is a section heading
            if DocumentProcessor._is_section_heading(line):
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content).strip()

                # Start new section
                current_section = line.strip()
                current_content = []
            else:
                current_content.append(line)

        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content).strip()

        logger.info(f"Extracted {len(sections)} sections")
        return sections

    @staticmethod
    def _is_section_heading(line: str) -> bool:
        """Check if line is likely a section heading."""
        line = line.strip()

        if not line:
            return False

        # Common patterns for section headings
        patterns = [
            line.startswith('SECTION'),
            line.startswith('PART'),
            line.startswith('CHAPTER'),
            line.endswith('===') or line.startswith('==='),
            line.endswith('---') or line.startswith('---'),
            line.isupper() and len(line.split()) <= 10,
            line.startswith('##'),  # Markdown style
        ]

        # Check for numbered sections (e.g., "1. Introduction", "1.1 Background")
        if line[0].isdigit() and ('.' in line[:10]):
            return True

        return any(patterns)

    @staticmethod
    def replace_mentions_in_document(
        original_path: Path,
        mention_replacements: Dict[str, str],
        output_path: Path
    ) -> bool:
        """
        Replace @mentions in document with actual values.

        Args:
            original_path: Original document path
            mention_replacements: Dict of mention -> replacement value
            output_path: Output file path

        Returns:
            True if successful

        Example:
            >>> replacements = {
            ...     '@companyname': 'Longboardfella Consulting Pty Ltd',
            ...     '@abn': '44 650 470 474'
            ... }
            >>> DocumentProcessor.replace_mentions_in_document(
            ...     original_path=Path('tender.docx'),
            ...     mention_replacements=replacements,
            ...     output_path=Path('tender_filled.docx')
            ... )
        """
        try:
            # Read original document
            text = DocumentProcessor.process_document(original_path)

            # Replace mentions
            for mention, value in mention_replacements.items():
                text = text.replace(mention, value)

            # Save to output
            output_format = output_path.suffix.lower().replace('.', '')
            return DocumentProcessor.save_document_with_mentions(
                text,
                output_path,
                output_format
            )

        except Exception as e:
            logger.error(f"Failed to replace mentions: {e}")
            return False
