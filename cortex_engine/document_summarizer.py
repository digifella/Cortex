# ## File: cortex_engine/document_summarizer.py
# Version: 1.0.0
# Date: 2025-08-28
# Purpose: Advanced document summarization engine for Cortex Suite.
#          Leverages existing document processing and LLM infrastructure.

import os
import json
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
import re

from .utils import get_logger
from .config import get_model_config_for_task, PROPOSAL_LLM_MODEL
from .docling_reader import DoclingDocumentReader, DOCLING_AVAILABLE

# Set up logging
logger = get_logger(__name__)

@dataclass
class SummaryResult:
    """Container for summarization results"""
    success: bool
    summary: str
    metadata: Dict[str, Any]
    error: Optional[str] = None
    processing_time: float = 0.0
    word_count: int = 0
    page_count: int = 0

class DocumentSummarizer:
    """
    Advanced document summarization engine using Cortex Suite's LLM infrastructure.
    
    Features:
    - Multiple summary levels (Highlights, Summary, Detailed)
    - Smart document chunking for large files
    - Integration with Docling for superior text extraction
    - Markdown-formatted output
    - Progress tracking for long operations
    """
    
    def __init__(self):
        """Initialize the document summarizer."""
        # Use the same model as proposals for consistency
        self.model_name = PROPOSAL_LLM_MODEL
        
        # Initialize document reader
        if DOCLING_AVAILABLE:
            self.document_reader = DoclingDocumentReader(
                ocr_enabled=True,
                table_structure_recognition=True
            )
            logger.info("✅ Using Docling for document processing")
        else:
            self.document_reader = None
            logger.warning("⚠️ Docling not available, using fallback text extraction")
    
    def _call_llm(self, prompt: str, max_tokens: int = 4000, temperature: float = 0.3) -> str:
        """Call LLM using Ollama API with GPU optimization."""
        import requests
        import json
        
        try:
            # First check if Ollama is accessible
            try:
                health_response = requests.get("http://localhost:11434/api/tags", timeout=5)
                health_response.raise_for_status()
                logger.info("✅ Ollama service is accessible")
            except requests.exceptions.RequestException as health_error:
                raise Exception(f"Ollama service not accessible: {health_error}")
            
            # Pre-load model to GPU if not already loaded (may take 30-60 seconds first time)
            self._ensure_model_loaded()
            
            url = "http://localhost:11434/api/generate"
            payload = {
                "model": self.model_name,
                "prompt": prompt,
                "stream": False,
                "keep_alive": "10m",  # Keep model in GPU memory for 10 minutes
                "options": {
                    "temperature": temperature,
                    "num_predict": max_tokens,
                    "num_gpu": -1,  # Use all available GPU layers
                    "num_thread": 8  # Optimize CPU threads when needed
                }
            }
            
            # Adaptive timeout based on prompt length and max_tokens
            base_timeout = 300  # 5 minutes base
            prompt_factor = len(prompt) // 10000  # +1 minute per 10k chars
            token_factor = max_tokens // 1000    # +1 minute per 1k tokens
            adaptive_timeout = min(base_timeout + prompt_factor * 60 + token_factor * 30, 900)  # Max 15 minutes
            
            logger.info(f"🔄 Calling Ollama with model: {self.model_name} (GPU accelerated)")
            logger.info(f"📊 Prompt: {len(prompt)} chars, Max tokens: {max_tokens}, Timeout: {adaptive_timeout//60}m{adaptive_timeout%60}s")
            response = requests.post(url, json=payload, timeout=adaptive_timeout)
            
            if response.status_code == 403:
                raise Exception("Ollama returned 403 Forbidden - check if the model is available and Ollama has proper permissions")
            elif response.status_code == 404:
                raise Exception(f"Model '{self.model_name}' not found in Ollama. Available models: {self._get_available_models()}")
            
            response.raise_for_status()
            
            result = response.json()
            
            if 'response' in result:
                logger.info("✅ Successfully received response from Ollama")
                return result['response']
            elif 'error' in result:
                raise Exception(f"Ollama returned error: {result['error']}")
            else:
                raise Exception(f"Unexpected response format: {result}")
                
        except requests.exceptions.Timeout:
            timeout_mins = adaptive_timeout // 60
            raise Exception(f"Ollama request timed out after {timeout_mins} minutes - try using 'highlights' level for large documents or check GPU memory")
        except requests.exceptions.ConnectionError:
            raise Exception("Cannot connect to Ollama service - ensure Ollama is running on localhost:11434")
        except requests.exceptions.RequestException as e:
            raise Exception(f"Ollama API request failed: {e}")
        except json.JSONDecodeError as e:
            raise Exception(f"Invalid JSON response from Ollama: {e}")
        except Exception as e:
            raise Exception(f"LLM call failed: {e}")
    
    def _get_available_models(self) -> str:
        """Get list of available models for error reporting."""
        try:
            import requests
            response = requests.get("http://localhost:11434/api/tags", timeout=5)
            response.raise_for_status()
            models = response.json().get('models', [])
            return ", ".join([model['name'] for model in models])
        except:
            return "Unable to fetch model list"
    
    def _ensure_model_loaded(self):
        """Ensure the model is loaded into GPU memory."""
        import requests
        
        try:
            # Check if model is already loaded
            ps_response = requests.get("http://localhost:11434/api/ps", timeout=5)
            if ps_response.status_code == 200:
                loaded_models = ps_response.json().get('models', [])
                for model in loaded_models:
                    if model.get('name', '').startswith(self.model_name):
                        logger.info(f"✅ Model {self.model_name} already loaded in GPU")
                        return
            
            # Pre-load model with a small prompt to get it into GPU memory
            logger.info(f"🔄 Pre-loading {self.model_name} into GPU memory...")
            preload_payload = {
                "model": self.model_name,
                "prompt": "Hello",
                "stream": False,
                "keep_alive": "10m",
                "options": {
                    "num_predict": 1,
                    "num_gpu": -1
                }
            }
            
            preload_response = requests.post(
                "http://localhost:11434/api/generate", 
                json=preload_payload, 
                timeout=60
            )
            
            if preload_response.status_code == 200:
                logger.info(f"✅ Model {self.model_name} loaded into GPU memory")
            else:
                logger.warning(f"⚠️ Failed to pre-load model: HTTP {preload_response.status_code}")
                
        except Exception as e:
            logger.warning(f"⚠️ Could not pre-load model: {e}")
            # Continue anyway - the main request might still work
    
    def summarize_document(
        self, 
        file_path: str, 
        summary_level: str = "summary",
        progress_callback: Optional[callable] = None
    ) -> SummaryResult:
        """
        Summarize a document with the specified detail level.
        
        Args:
            file_path: Path to the document to summarize
            summary_level: One of 'highlights', 'summary', 'detailed'
            progress_callback: Optional callback for progress updates
            
        Returns:
            SummaryResult with the summary and metadata
        """
        import time
        start_time = time.time()
        
        try:
            if progress_callback:
                progress_callback("Reading document...", 5)
            
            # Extract text from document
            document_content, metadata = self._extract_document_content(file_path)
            
            if not document_content.strip():
                return SummaryResult(
                    success=False,
                    summary="",
                    metadata={},
                    error="Could not extract text from document"
                )
            
            if progress_callback:
                progress_callback(f"Extracted {len(document_content)} characters from document", 15)
            
            # Get word count and estimate pages
            word_count = len(document_content.split())
            estimated_pages = max(1, word_count // 250)  # ~250 words per page
            
            if progress_callback:
                progress_callback(f"Analyzing {word_count:,} words (~{estimated_pages} pages)", 25)
            
            # Check if document needs chunking - be more aggressive for "summary" and "detailed" levels
            if summary_level == "highlights":
                max_tokens = 15000  # Smaller chunks for highlights
                chunk_overlap = 500
            elif summary_level == "summary":
                max_tokens = 8000   # Smaller chunks for better processing
                chunk_overlap = 1000
            else:  # detailed
                max_tokens = 6000   # Even smaller for detailed analysis
                chunk_overlap = 1500
            
            estimated_tokens = word_count * 1.3  # Rough token estimation
            
            if estimated_tokens > max_tokens or summary_level in ["summary", "detailed"]:
                if progress_callback:
                    progress_callback(f"Using intelligent chunking for {summary_level} analysis", 30)
                summary = self._summarize_large_document(document_content, summary_level, progress_callback, max_tokens, chunk_overlap)
            else:
                if progress_callback:
                    progress_callback(f"Sending to AI for {summary_level} analysis...", 40)
                summary = self._generate_summary(document_content, summary_level, metadata)
            
            if progress_callback:
                progress_callback("Finalizing...", 90)
            
            processing_time = time.time() - start_time
            
            return SummaryResult(
                success=True,
                summary=summary,
                metadata={
                    **metadata,
                    'summary_level': summary_level,
                    'processing_time_seconds': round(processing_time, 2)
                },
                processing_time=processing_time,
                word_count=word_count,
                page_count=estimated_pages
            )
            
        except Exception as e:
            logger.error(f"Document summarization failed: {e}")
            return SummaryResult(
                success=False,
                summary="",
                metadata={},
                error=str(e),
                processing_time=time.time() - start_time
            )
    
    def _extract_document_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text content from document using available readers."""
        file_path = Path(file_path)
        metadata = {
            'filename': file_path.name,
            'file_size_mb': round(file_path.stat().st_size / (1024 * 1024), 2),
            'file_extension': file_path.suffix.lower()
        }
        
        # Try Docling first if available
        if self.document_reader and metadata['file_extension'] in ['.pdf', '.docx', '.pptx', '.xlsx']:
            try:
                documents = self.document_reader.load_data([str(file_path)])
                if documents:
                    content = "\n\n".join([doc.text for doc in documents])
                    logger.info(f"✅ Extracted {len(content)} characters using Docling")
                    return content, metadata
            except Exception as e:
                logger.warning(f"Docling extraction failed, using fallback: {e}")
        
        # Fallback to simple text extraction
        try:
            if metadata['file_extension'] == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif metadata['file_extension'] == '.pdf':
                # Use PyMuPDF as fallback
                try:
                    import fitz  # PyMuPDF
                    doc = fitz.open(str(file_path))
                    content = ""
                    for page in doc:
                        content += page.get_text()
                    doc.close()
                except ImportError:
                    raise Exception("PDF processing requires PyMuPDF (pip install PyMuPDF)")
            elif metadata['file_extension'] == '.docx':
                # Use python-docx as fallback
                try:
                    import docx
                    doc = docx.Document(file_path)
                    content = ""
                    for paragraph in doc.paragraphs:
                        content += paragraph.text + "\n"
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                content += cell.text + " "
                        content += "\n"
                except ImportError:
                    raise Exception("DOCX processing requires python-docx (pip install python-docx)")
            elif metadata['file_extension'] in ['.pptx', '.xlsx']:
                # For now, suggest using Docling for these formats
                raise Exception(f"File type {metadata['file_extension']} requires Docling for processing. Please install with: pip install docling")
            else:
                raise Exception(f"Unsupported file type: {metadata['file_extension']}")
            
            logger.info(f"✅ Extracted {len(content)} characters using fallback reader")
            return content, metadata
            
        except Exception as e:
            logger.error(f"Document extraction failed: {e}")
            raise
    
    def _generate_summary(self, content: str, summary_level: str, metadata: Dict[str, Any]) -> str:
        """Generate summary using LLM."""
        
        # Create prompt based on summary level
        prompt = self._create_summary_prompt(content, summary_level, metadata.get('filename', 'document'))
        
        try:
            # Use Ollama directly for summarization
            response = self._call_llm(
                prompt=prompt,
                max_tokens=2000 if summary_level == 'highlights' else 4000,
                temperature=0.3  # Lower temperature for more focused summaries
            )
            
            return response.strip()
                
        except Exception as e:
            logger.error(f"Summary generation failed: {e}")
            raise
    
    def _create_summary_prompt(self, content: str, summary_level: str, filename: str) -> str:
        """Create appropriate prompt based on summary level."""
        
        base_instructions = f"""You are an expert document analyst. Analyze the following document ({filename}) and provide a {summary_level} summary.

Key requirements:
- Output ONLY in clean, well-structured Markdown format
- Be accurate and factual - do not invent information
- Focus on the most important and actionable content
- Use proper heading hierarchy (##, ###, etc.)
- Include bullet points and lists where appropriate
- Maintain professional tone throughout"""

        if summary_level == 'highlights':
            specific_instructions = """
**HIGHLIGHTS ONLY FORMAT** (2-3 key points maximum):
- Provide a 2-3 sentence executive summary
- List 2-3 most critical takeaways as bullet points
- Include key numbers, dates, or decisions if present
- Keep total output under 150 words

Structure your response as:
## 📋 Document Highlights
[Brief summary paragraph]

**Key Takeaways:**
- [Most important point]
- [Second most important point]  
- [Third point if applicable]
"""
        elif summary_level == 'summary':
            specific_instructions = """
**SUMMARY FORMAT** (1-2 paragraphs per major section):
- Document purpose and context
- Main topics and sections covered
- Key findings, recommendations, or conclusions
- Important data points and evidence
- Actionable insights for the reader
- Keep total output between 300-600 words

Structure your response as:
## 📄 Document Summary

### Purpose & Context
[What this document is about and why it was created]

### Main Topics
[Key areas covered in the document]

### Key Findings
[Important conclusions, recommendations, or results]

### Actionable Insights
[What the reader should do with this information]
"""
        else:  # detailed
            specific_instructions = """
**DETAILED SUMMARY FORMAT** (Comprehensive analysis):
- Section-by-section breakdown of the document
- Detailed findings and analysis
- Supporting evidence and examples
- Implementation recommendations
- Risk factors and considerations
- Background context and methodology
- Comprehensive insights and implications

Structure your response as:
## 📖 Detailed Document Analysis

### Executive Summary
[High-level overview of the entire document]

### Document Structure & Purpose
[How the document is organized and its intended purpose]

### Detailed Findings
[Section-by-section analysis with supporting details]

### Key Evidence & Examples
[Important data, case studies, or supporting information]

### Recommendations & Next Steps
[What should be done based on this document]

### Risks & Considerations
[Potential challenges or limitations to consider]

### Implementation Guidance
[How to act on the information provided]
"""

        return f"""{base_instructions}

{specific_instructions}

DOCUMENT CONTENT:
---
{content[:20000]}  # Limit content to avoid token limits
---

Please analyze the above document and provide your {summary_level} summary in the requested Markdown format:"""

    def _summarize_large_document(self, content: str, summary_level: str, progress_callback: Optional[callable] = None, max_chunk_tokens: int = 8000, chunk_overlap: int = 1000) -> str:
        """Handle large documents by chunking and merging summaries."""
        
        # Split content into logical chunks (by sections/paragraphs)
        chunks = self._intelligent_chunk_document(content, max_chunk_tokens)
        logger.info(f"Split document into {len(chunks)} chunks for processing")
        
        chunk_summaries = []
        
        for i, chunk in enumerate(chunks):
            if progress_callback:
                progress = 40 + (40 * (i + 1) / len(chunks))  # 40-80% range
                progress_callback(f"Processing chunk {i+1}/{len(chunks)}...", progress)
            
            try:
                # For chunked processing, always use 'highlights' for individual chunks to avoid timeouts
                chunk_level = 'highlights' if summary_level in ['summary', 'detailed'] else summary_level
                chunk_summary = self._generate_summary(chunk, chunk_level, {'filename': f'chunk_{i+1}'})
                chunk_summaries.append(chunk_summary)
            except Exception as e:
                logger.warning(f"Failed to summarize chunk {i+1}: {e}")
                # Continue with other chunks, but add a placeholder
                chunk_summaries.append(f"[Chunk {i+1} processing failed - content may be too complex]")
        
        if not chunk_summaries:
            raise Exception("Failed to summarize any document chunks")
        
        # Merge chunk summaries into final summary
        if progress_callback:
            progress_callback("Consolidating results...", 85)
        
        return self._merge_chunk_summaries(chunk_summaries, summary_level)
    
    def _intelligent_chunk_document(self, content: str, max_chunk_tokens: int = 8000) -> List[str]:
        """Split document into chunks at logical boundaries."""
        
        # Estimate tokens (rough approximation)
        words = content.split()
        total_tokens = len(words) * 1.3
        
        if total_tokens <= max_chunk_tokens:
            return [content]
        
        # Try to split by sections first
        section_patterns = [
            r'\n\s*#{1,6}\s+',  # Markdown headers
            r'\n\s*\d+\.\s+',   # Numbered sections
            r'\n\s*[A-Z][^a-z\n]{10,50}\n',  # All-caps section headers
            r'\n\s*\n\s*\n',    # Multiple line breaks
        ]
        
        chunks = [content]
        
        for pattern in section_patterns:
            new_chunks = []
            for chunk in chunks:
                if len(chunk.split()) * 1.3 > max_chunk_tokens:
                    # Split this chunk
                    parts = re.split(pattern, chunk)
                    if len(parts) > 1:
                        new_chunks.extend([part.strip() for part in parts if part.strip()])
                    else:
                        new_chunks.append(chunk)
                else:
                    new_chunks.append(chunk)
            chunks = new_chunks
        
        # Final fallback: split by paragraphs or sentences if still too large
        final_chunks = []
        for chunk in chunks:
            if len(chunk.split()) * 1.3 <= max_chunk_tokens:
                final_chunks.append(chunk)
            else:
                # Force split by paragraphs
                paragraphs = chunk.split('\n\n')
                current_chunk = ""
                
                for paragraph in paragraphs:
                    test_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph
                    if len(test_chunk.split()) * 1.3 <= max_chunk_tokens:
                        current_chunk = test_chunk
                    else:
                        if current_chunk:
                            final_chunks.append(current_chunk.strip())
                        current_chunk = paragraph
                
                if current_chunk:
                    final_chunks.append(current_chunk.strip())
        
        return [chunk for chunk in final_chunks if chunk.strip()]
    
    def _merge_chunk_summaries(self, chunk_summaries: List[str], target_level: str) -> str:
        """Merge individual chunk summaries into final summary."""
        
        combined_content = "\n\n---\n\n".join(chunk_summaries)
        
        merge_prompt = f"""You are an expert document analyst. I have processed a large document in chunks and need you to create a unified {target_level} summary from the individual chunk summaries.

Your task is to:
1. Synthesize the information from all chunks into a coherent whole
2. Remove redundancy and consolidate similar points
3. Maintain the most important information from across all chunks
4. Create a {target_level} summary that reads as if it came from analyzing the full document

FORMAT: Output in clean Markdown format appropriate for the {target_level} level.

CHUNK SUMMARIES TO MERGE:
---
{combined_content}
---

Please provide the unified {target_level} summary:"""

        try:
            response = self._call_llm(
                prompt=merge_prompt,
                max_tokens=3000 if target_level == 'detailed' else 2000,
                temperature=0.2  # Very focused for merging
            )
            
            return response.strip()
                
        except Exception as e:
            logger.error(f"Summary merge failed: {e}")
            return self._fallback_merge(chunk_summaries, target_level)
    
    def _fallback_merge(self, chunk_summaries: List[str], target_level: str) -> str:
        """Fallback merge method if LLM merge fails."""
        
        level_titles = {
            'highlights': '📋 Document Highlights',
            'summary': '📄 Document Summary', 
            'detailed': '📖 Detailed Document Analysis'
        }
        
        result = f"## {level_titles.get(target_level, 'Document Summary')}\n\n"
        result += "*Note: This document was processed in multiple sections*\n\n"
        
        for i, summary in enumerate(chunk_summaries, 1):
            result += f"### Section {i}\n\n{summary}\n\n---\n\n"
        
        return result.strip()