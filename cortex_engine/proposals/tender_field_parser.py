"""
Tender Field Parser - Detect Fillable Fields in Tender Documents

This module detects all fillable fields in tender documents with priority on:
1. Table-based fields (most common in standard tenders)
2. Placeholder text ([INSERT], ___, TBD, etc.)
3. Question-based sections

Author: Cortex Suite
Created: 2026-01-04
Version: 1.0.0
"""

import re
import docx
from typing import List, Dict, Optional, Tuple, Any, TYPE_CHECKING
from dataclasses import dataclass, field as dataclass_field
from enum import Enum
import logging

# Handle imports for both module and direct execution
if TYPE_CHECKING or __name__ != "__main__":
    from .field_classifier import FieldClassifier, get_classifier
else:
    # Will be imported in __main__ block for direct execution
    FieldClassifier = None
    get_classifier = None

logger = logging.getLogger(__name__)


class LocationType(str, Enum):
    """Type of location where field was found."""
    TABLE_CELL = "table_cell"
    PARAGRAPH = "paragraph"
    HEADING = "heading"
    LIST_ITEM = "list_item"


@dataclass
class DetectedField:
    """
    Represents a detected fillable field in a tender document.

    This stores comprehensive information about where the field was found,
    what type of information is needed, and how to fill it.
    """
    # Identification
    field_id: str                           # Unique ID (e.g., "table_2_r3_c1", "para_45")

    # Location information
    location: str                           # Human-readable location
    location_type: LocationType             # Type of location
    location_coordinates: Dict[str, Any]    # Precise coordinates for filling

    # Field description
    field_description: str                  # Question/label text
    field_context: str                      # Surrounding text for context

    # Classification
    classification_hints: List[str] = dataclass_field(default_factory=list)  # Suggested field types
    classification_confidence: float = 0.0  # Confidence in classification
    classification_method: str = "none"     # How it was classified

    # Current state
    is_empty: bool = True                   # Whether field is currently empty
    current_value: Optional[str] = None     # Current value (if any)

    # Metadata
    table_context: Optional[Dict] = None    # Additional table-specific context


class TenderFieldParser:
    """
    Detects all fillable fields in tender documents.

    Priority order:
    1. Table cells (most common in standard tenders)
    2. Placeholder text patterns
    3. Question-based sections
    """

    # Reuse placeholder patterns from FlexibleTemplateParser
    PLACEHOLDER_PATTERNS = [
        r'\[.*?\]',                          # [INSERT], [COMPANY NAME]
        r'<.*?>',                            # <insert text>
        r'_{3,}',                            # _____ (underscores)
        r'\.{3,}',                           # ..... (dots)
        r'\bTBD\b',                          # TBD
        r'\bTBC\b',                          # TBC (To Be Confirmed)
        r'\bXXX\b',                          # XXX
        r'\bTO BE COMPLETED\b',              # TO BE COMPLETED
        r'\bPLEASE COMPLETE\b',              # PLEASE COMPLETE
        r'\bINSERT\b',                       # INSERT
        r'\bTO BE PROVIDED\b',               # TO BE PROVIDED
    ]

    # Patterns for question detection
    QUESTION_PATTERNS = [
        r'^(?:what|how|why|when|where|who|which|describe|explain|provide|list|detail)',
        r'\?$',  # Ends with question mark
    ]

    def __init__(
        self,
        use_classification: bool = True,
        classifier: Optional[FieldClassifier] = None
    ):
        """
        Initialize the tender field parser.

        Args:
            use_classification: Whether to auto-classify fields
            classifier: Optional FieldClassifier instance (creates one if None)
        """
        self.use_classification = use_classification
        self.classifier = classifier or get_classifier() if use_classification else None

        # Compile patterns for performance
        self._compiled_placeholders = [
            re.compile(pattern, re.IGNORECASE)
            for pattern in self.PLACEHOLDER_PATTERNS
        ]
        self._compiled_questions = [
            re.compile(pattern, re.IGNORECASE)
            for pattern in self.QUESTION_PATTERNS
        ]

        logger.info("TenderFieldParser initialized")

    def parse_tender_document(self, doc: docx.Document) -> List[DetectedField]:
        """
        Main entry point - parse entire tender document.

        Args:
            doc: python-docx Document object

        Returns:
            List of DetectedField objects
        """
        logger.info("Starting tender document parsing...")

        detected_fields = []

        # Priority 1: Parse tables (most common in standard tenders)
        table_fields = self._parse_tables(doc)
        detected_fields.extend(table_fields)
        logger.info(f"Found {len(table_fields)} fields in tables")

        # Priority 2: Parse paragraphs for placeholders and questions
        paragraph_fields = self._parse_paragraphs(doc)
        detected_fields.extend(paragraph_fields)
        logger.info(f"Found {len(paragraph_fields)} fields in paragraphs")

        logger.info(f"Total fields detected: {len(detected_fields)}")

        return detected_fields

    def _parse_tables(self, doc: docx.Document) -> List[DetectedField]:
        """
        Parse all tables in the document to find fillable cells.

        Strategy:
        - Iterate through all tables
        - For each cell, check if empty or contains placeholder
        - Extract row/column headers for context
        - Create DetectedField with precise coordinates
        """
        fields = []

        for table_idx, table in enumerate(doc.tables):
            logger.debug(f"Parsing table {table_idx + 1}")

            # Try to detect headers (first row or first column)
            header_row = self._extract_table_headers(table)

            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    # Skip merged cells that we've already processed
                    if self._is_merged_cell(cell, row_idx, col_idx, table):
                        continue

                    cell_text = cell.text.strip()

                    # Check if this cell needs filling
                    is_empty = len(cell_text) == 0
                    has_placeholder = self._has_placeholder(cell_text)

                    if is_empty or has_placeholder:
                        # Extract context from headers
                        field_description, context = self._extract_table_field_context(
                            table, row_idx, col_idx, header_row, cell_text
                        )

                        # Classify the field
                        classification_hints = []
                        classification_confidence = 0.0
                        classification_method = "none"

                        if self.use_classification and field_description:
                            result = self.classifier.classify(field_description)
                            classification_hints = result.field_types
                            classification_confidence = result.confidence
                            classification_method = result.classification_method

                        # Create field ID
                        field_id = f"table_{table_idx}_r{row_idx}_c{col_idx}"

                        # Create DetectedField
                        detected_field = DetectedField(
                            field_id=field_id,
                            location=f"Table {table_idx + 1}, Row {row_idx + 1}, Column {col_idx + 1}",
                            location_type=LocationType.TABLE_CELL,
                            location_coordinates={
                                "table_idx": table_idx,
                                "row_idx": row_idx,
                                "col_idx": col_idx,
                            },
                            field_description=field_description,
                            field_context=context,
                            classification_hints=classification_hints,
                            classification_confidence=classification_confidence,
                            classification_method=classification_method,
                            is_empty=is_empty,
                            current_value=cell_text if not is_empty else None,
                            table_context={
                                "header_row": header_row,
                                "row_header": self._get_row_header(table, row_idx),
                                "column_header": header_row[col_idx] if col_idx < len(header_row) else None,
                            }
                        )

                        fields.append(detected_field)
                        logger.debug(f"Detected field: {field_id} - {field_description}")

        return fields

    def _parse_paragraphs(self, doc: docx.Document) -> List[DetectedField]:
        """
        Parse paragraphs for placeholder text and questions.

        Strategy:
        - Check each paragraph for placeholders or question patterns
        - Extract context from surrounding paragraphs
        - Handle numbered/bulleted lists
        """
        fields = []

        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()

            if not text:
                continue

            # Check for placeholders
            has_placeholder = self._has_placeholder(text)

            # Check for questions
            is_question = self._is_question(text)

            if has_placeholder or is_question:
                # Extract context from previous paragraph if needed
                context = self._extract_paragraph_context(doc, para_idx)

                # Use the text itself as field description
                field_description = text

                # Classify the field
                classification_hints = []
                classification_confidence = 0.0
                classification_method = "none"

                if self.use_classification:
                    result = self.classifier.classify(field_description)
                    classification_hints = result.field_types
                    classification_confidence = result.confidence
                    classification_method = result.classification_method

                # Create field ID
                field_id = f"para_{para_idx}"

                # Create DetectedField
                detected_field = DetectedField(
                    field_id=field_id,
                    location=f"Paragraph {para_idx + 1}",
                    location_type=LocationType.PARAGRAPH,
                    location_coordinates={
                        "paragraph_idx": para_idx,
                    },
                    field_description=field_description,
                    field_context=context,
                    classification_hints=classification_hints,
                    classification_confidence=classification_confidence,
                    classification_method=classification_method,
                    is_empty=has_placeholder,
                    current_value=text if not has_placeholder else None,
                )

                fields.append(detected_field)
                logger.debug(f"Detected field: {field_id} - {field_description[:50]}...")

        return fields

    # ==================== Helper Methods ====================

    def _has_placeholder(self, text: str) -> bool:
        """Check if text contains placeholder patterns."""
        if not text:
            return False
        return any(pattern.search(text) for pattern in self._compiled_placeholders)

    def _is_question(self, text: str) -> bool:
        """Check if text appears to be a question."""
        if not text:
            return False
        return any(pattern.search(text) for pattern in self._compiled_questions)

    def _extract_table_headers(self, table: docx.table.Table) -> List[str]:
        """
        Extract header row from table (assumes first row is headers).

        Returns:
            List of header texts
        """
        if not table.rows:
            return []

        header_row = table.rows[0]
        headers = [cell.text.strip() for cell in header_row.cells]

        return headers

    def _get_row_header(self, table: docx.table.Table, row_idx: int) -> Optional[str]:
        """Get the first cell of a row (often used as row label)."""
        if row_idx >= len(table.rows):
            return None

        row = table.rows[row_idx]
        if row.cells:
            return row.cells[0].text.strip()

        return None

    def _extract_table_field_context(
        self,
        table: docx.table.Table,
        row_idx: int,
        col_idx: int,
        header_row: List[str],
        cell_text: str
    ) -> Tuple[str, str]:
        """
        Extract field description and context from table cell.

        Returns:
            (field_description, context)
        """
        # Build field description from headers
        description_parts = []

        # Add column header
        if col_idx < len(header_row) and header_row[col_idx]:
            description_parts.append(header_row[col_idx])

        # Add row header (first column of this row)
        row_header = self._get_row_header(table, row_idx)
        if row_header and col_idx > 0:  # Don't duplicate if we're in the header column
            description_parts.append(f"for {row_header}")

        # If we have placeholder text, include it
        if cell_text:
            description_parts.append(f"({cell_text})")

        field_description = " ".join(description_parts) if description_parts else f"Table {row_idx},{col_idx}"

        # Context is the surrounding headers
        context_parts = []
        if header_row:
            context_parts.append(f"Columns: {', '.join(header_row)}")
        if row_header:
            context_parts.append(f"Row: {row_header}")

        context = " | ".join(context_parts)

        return field_description, context

    def _extract_paragraph_context(
        self,
        doc: docx.Document,
        para_idx: int,
        context_window: int = 2
    ) -> str:
        """
        Extract context from surrounding paragraphs.

        Args:
            doc: Document object
            para_idx: Index of current paragraph
            context_window: Number of paragraphs before/after to include

        Returns:
            Context string
        """
        context_parts = []

        # Get previous paragraphs for context
        start_idx = max(0, para_idx - context_window)
        for i in range(start_idx, para_idx):
            prev_text = doc.paragraphs[i].text.strip()
            if prev_text:
                context_parts.append(prev_text)

        return " | ".join(context_parts[-2:])  # Last 2 non-empty paragraphs

    def _is_merged_cell(
        self,
        cell: docx.table._Cell,
        row_idx: int,
        col_idx: int,
        table: docx.table.Table
    ) -> bool:
        """
        Check if this is a merged cell that we've already processed.

        Note: This is a simplified check. Full merged cell detection is complex in python-docx.
        """
        # For now, simple heuristic: check if this cell is identical to previous cell
        # This catches horizontally merged cells
        if col_idx > 0:
            prev_cell = table.rows[row_idx].cells[col_idx - 1]
            if cell == prev_cell:  # Same cell object = merged
                return True

        return False


if __name__ == "__main__":
    # Quick test with a sample document
    import sys
    from pathlib import Path
    sys.path.insert(0, str(Path(__file__).parent.parent.parent))

    from cortex_engine.proposals.field_classifier import FieldClassifier, get_classifier

    logging.basicConfig(level=logging.INFO)

    print("=" * 80)
    print("TENDER FIELD PARSER TEST")
    print("=" * 80)

    # Create a simple test document
    test_doc = docx.Document()

    # Add a table with some fillable fields
    test_doc.add_heading("Tenderer Information", level=1)
    table = test_doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Response"
    table.cell(1, 0).text = "ABN"
    table.cell(1, 1).text = ""  # Empty - should be detected
    table.cell(2, 0).text = "Email"
    table.cell(2, 1).text = "[INSERT EMAIL]"  # Placeholder - should be detected
    table.cell(3, 0).text = "Experience"
    table.cell(3, 1).text = "Please describe your relevant experience"  # Question

    # Add some paragraph-based fields
    test_doc.add_paragraph("Describe your approach to quality management: _______________")
    test_doc.add_paragraph("What makes your organization uniquely qualified for this project?")

    # Parse the document
    parser = TenderFieldParser(use_classification=True)
    fields = parser.parse_tender_document(test_doc)

    print(f"\n✅ Detected {len(fields)} fillable fields:\n")

    for field in fields:
        print(f"ID: {field.field_id}")
        print(f"Location: {field.location}")
        print(f"Description: {field.field_description}")
        print(f"Classification: {field.classification_hints} (confidence: {field.classification_confidence:.2f}, method: {field.classification_method})")
        print(f"Empty: {field.is_empty}")
        print("-" * 80)
