from typing import List, Dict
from dataclasses import dataclass
import io
import docx


@dataclass
class _Instruction:
    section_heading: str
    instruction_raw: str
    task_type: str
    parameter: str = ""
    placeholder_paragraph: str = ""
    sub_instruction: str = ""


class TaskExecutionEngine:
    """Minimal TaskExecutionEngine for Docker distribution.
    Provides stubbed generation/refinement/retrieval and assembles a basic docx.
    """

    def __init__(self, main_index=None, collection_manager=None):
        self.index = main_index
        self.collection_manager = collection_manager

    def refine_with_ai(self, section_heading: str, raw_text: str, creativity: str = "green", sub_instruction: str = "") -> str:
        prefix = {
            "green": "[Factual]",
            "orange": "[Persuasive]",
            "red": "[Visionary]",
        }.get(creativity, "[Factual]")
        hint = f" Hint: {sub_instruction}" if sub_instruction else ""
        return f"{prefix} {raw_text}{hint}"

    def generate_from_kb(self, instruction, creativity: str, knowledge_sources: List[str], session_state) -> str:
        heading = getattr(instruction, 'section_heading', 'Section')
        task = getattr(instruction, 'task_type', 'GENERATE_FROM_KB')
        hint = getattr(instruction, 'sub_instruction', '')
        hint_txt = f" Hint: {hint}" if hint else ""
        return f"Generated draft for '{heading}' using sources {knowledge_sources} ({task}).{hint_txt}"

    def retrieve_from_kb(self, instruction, knowledge_sources: List[str]) -> str:
        heading = getattr(instruction, 'section_heading', 'Section')
        return f"Key points retrieved for '{heading}' from {knowledge_sources}:\n- Point A\n- Point B\n- Point C"

    def assemble_document(self, parsed_instructions: List, section_content: Dict, template_bytes: bytes) -> bytes:
        # Create a simple doc using headings and content from section_content in order
        doc = docx.Document(io.BytesIO(template_bytes)) if template_bytes else docx.Document()
        for i, inst in enumerate(parsed_instructions):
            heading = getattr(inst, 'section_heading', f'Section {i+1}')
            content_key = f"content_inst_{i}"
            content = section_content.get(content_key, {}).get('text', '')
            if heading:
                doc.add_heading(heading, level=2)
            doc.add_paragraph(content or "[No content]")
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()

