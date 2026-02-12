"""
Proposal Export Engine
Version: 1.0.0
Date: 2026-02-01

Purpose: Analyze completeness and export completed proposals as Markdown or DOCX.
Merges Tier 1 (entity profile fields) and Tier 2 (AI-generated responses) into
the original tender document with all substitutions applied.
"""

import re
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Any, Tuple
from datetime import datetime
from pathlib import Path

from .workspace_manager import WorkspaceManager
from .workspace_model import Workspace, MentionBinding
from .entity_profile_manager import EntityProfileManager
from .entity_profile_schema import EntityProfile
from .field_classifier import FieldTier, ClassifiedField
from .ic_persistence_model import (
    PersistedClassifiedField,
    persisted_to_classified_field
)
from .utils import get_logger

logger = get_logger(__name__)


@dataclass
class CompletionItem:
    """A single item that needs completion."""
    field_text: str
    tier: str  # "tier1" or "tier2"
    status: str  # "completed", "pending", "missing"
    resolved_value: Optional[str] = None
    confidence: Optional[float] = None
    has_placeholders: bool = False


@dataclass
class CompletionReport:
    """Report on proposal completion status."""
    total_fields: int = 0
    completed_fields: int = 0
    missing_fields: int = 0
    tier1_total: int = 0
    tier1_completed: int = 0
    tier2_total: int = 0
    tier2_completed: int = 0
    completion_percentage: float = 0.0
    quality_flags: List[str] = field(default_factory=list)
    items: List[CompletionItem] = field(default_factory=list)

    @property
    def tier1_percentage(self) -> float:
        return (self.tier1_completed / self.tier1_total * 100) if self.tier1_total > 0 else 0.0

    @property
    def tier2_percentage(self) -> float:
        return (self.tier2_completed / self.tier2_total * 100) if self.tier2_total > 0 else 0.0


class ProposalExportEngine:
    """Analyzes completeness and exports completed proposals."""

    # Mapping from field_classifier profile_field to entity profile attribute paths
    PROFILE_FIELD_MAP = {
        "company_name": lambda e: e.company.legal_name,
        "trading_name": lambda e: e.company.trading_names[0] if e.company.trading_names else "",
        "abn": lambda e: e.format_abn() or (e.company.abn or ""),
        "acn": lambda e: e.format_acn() or (e.company.acn or ""),
        "arbn": lambda e: "",  # Not in schema
        "registered_address": lambda e: e.contact.registered_office.formatted(single_line=True),
        "postal_address": lambda e: (e.contact.postal_address.formatted(single_line=True) if e.contact.postal_address else e.contact.registered_office.formatted(single_line=True)),
        "business_address": lambda e: e.contact.registered_office.formatted(single_line=True),
        "contact_email": lambda e: e.contact.email,
        "contact_phone": lambda e: e.contact.phone,
        "contact_mobile": lambda e: "",  # Not a separate field
        "contact_fax": lambda e: "",  # Not in schema
        "website": lambda e: e.contact.website or "",
        "contact_person_name": lambda e: "",  # Would need specific contact
        "contact_person_title": lambda e: "",
        "bank_name": lambda e: "",
        "bank_bsb": lambda e: "",
        "bank_account_number": lambda e: "",
        "insurance_policy_number": lambda e: "",
        "insurance_details": lambda e: "",
        "incorporation_date": lambda e: str(e.company.registration_date) if e.company.registration_date else "",
    }

    def __init__(
        self,
        workspace_manager: WorkspaceManager,
        entity_manager: EntityProfileManager
    ):
        self.workspace_manager = workspace_manager
        self.entity_manager = entity_manager

    def resolve_profile_field(self, entity: EntityProfile, profile_field: str) -> Optional[str]:
        """Resolve a profile field name to a value from the entity profile."""
        # Check custom fields first
        if profile_field.startswith("custom_fields."):
            cf_name = profile_field.replace("custom_fields.", "")
            val = entity.get_custom_field_value(cf_name)
            return val

        # Check mapped fields
        resolver = self.PROFILE_FIELD_MAP.get(profile_field)
        if resolver:
            try:
                val = resolver(entity)
                return val if val else None
            except Exception:
                return None

        return None

    def analyze_completeness(
        self,
        workspace_id: str,
    ) -> CompletionReport:
        """
        Analyze how complete a proposal workspace is.

        Checks:
        - Tier 1 mention bindings (approved with resolved values)
        - Tier 2 IC question responses (completed status)
        - Quality flags (placeholders, low confidence, missing values)
        """
        workspace = self.workspace_manager.get_workspace(workspace_id)
        if not workspace:
            return CompletionReport()

        report = CompletionReport()
        entity = None
        if workspace.metadata.entity_id:
            entity = self.entity_manager.get_entity_profile(workspace.metadata.entity_id)

        # --- Tier 1: Mention bindings ---
        for mention in workspace.mentions:
            item = CompletionItem(
                field_text=mention.mention_text,
                tier="tier1",
                status="pending"
            )

            # Try to resolve value
            resolved = mention.resolved_value
            if not resolved and entity and mention.field_path:
                resolved = self.resolve_profile_field(entity, mention.field_path)

            if mention.approved and resolved:
                item.status = "completed"
                item.resolved_value = resolved
                report.tier1_completed += 1
            elif mention.rejected or mention.ignored:
                item.status = "completed"  # intentionally skipped
                report.tier1_completed += 1
            else:
                item.status = "pending"
                report.missing_fields += 1

            report.tier1_total += 1
            report.items.append(item)

        # --- Tier 2: IC completion state ---
        ic_state = self.workspace_manager.get_ic_completion_state(workspace_id)
        if ic_state:
            question_status = ic_state.get('question_status', {})
            for field_text, qs in question_status.items():
                status = qs.get('status', 'pending')
                response = qs.get('response', '')
                confidence = qs.get('confidence')

                item = CompletionItem(
                    field_text=field_text,
                    tier="tier2",
                    confidence=confidence
                )

                if status == 'completed' and response:
                    item.status = "completed"
                    item.resolved_value = response
                    report.tier2_completed += 1

                    # Check for placeholders
                    placeholder_patterns = [
                        r'\[NEEDS DETAIL:[^\]]+\]',
                        r'\[NEEDS VERIFICATION:[^\]]+\]',
                        r'\[CONFIRM:[^\]]+\]',
                        r'\[GENERATION FAILED:[^\]]+\]',
                    ]
                    for pattern in placeholder_patterns:
                        if re.search(pattern, response):
                            item.has_placeholders = True
                            break

                    if item.has_placeholders:
                        report.quality_flags.append(f"Placeholders in: {field_text[:50]}")

                    if confidence is not None and confidence < 0.5:
                        report.quality_flags.append(f"Low confidence ({confidence:.0%}): {field_text[:50]}")
                else:
                    item.status = "pending"
                    report.missing_fields += 1

                report.tier2_total += 1
                report.items.append(item)

        # Totals
        report.total_fields = report.tier1_total + report.tier2_total
        report.completed_fields = report.tier1_completed + report.tier2_completed
        report.missing_fields = report.total_fields - report.completed_fields
        if report.total_fields > 0:
            report.completion_percentage = (report.completed_fields / report.total_fields) * 100

        return report

    def build_completed_document(
        self,
        original_text: str,
        mentions: List[MentionBinding],
        ic_responses: Dict[str, str],
        entity: Optional[EntityProfile] = None
    ) -> str:
        """
        Merge all completions into the original document text.

        Args:
            original_text: The original tender document text
            mentions: Approved mention bindings with resolved values
            ic_responses: Dict of field_text -> response text from IC
            entity: Entity profile for resolving unresolved mentions

        Returns:
            Completed document text with all substitutions applied
        """
        result = original_text

        # Apply mention substitutions (Tier 1)
        for mention in mentions:
            if mention.rejected or mention.ignored:
                continue

            value = mention.resolved_value
            if not value and entity and mention.field_path:
                value = self.resolve_profile_field(entity, mention.field_path)

            if value:
                # Replace @mention_text patterns in document
                patterns = [
                    re.escape(mention.mention_text),
                    re.escape(f"@{mention.mention_text}"),
                ]
                for pattern in patterns:
                    result = re.sub(pattern, value, result, flags=re.IGNORECASE)

        return result

    def generate_export_markdown(
        self,
        workspace_id: str,
        include_citations: bool = False,
        flag_incomplete: bool = True
    ) -> str:
        """
        Generate a complete Markdown export of the proposal.

        Returns markdown text with:
        - Header with metadata
        - Tier 1 substitutions applied to document
        - Tier 2 responses inserted
        - Optional citations
        - Flags for incomplete items
        """
        workspace = self.workspace_manager.get_workspace(workspace_id)
        if not workspace:
            return "# Error: Workspace not found"

        entity = None
        if workspace.metadata.entity_id:
            entity = self.entity_manager.get_entity_profile(workspace.metadata.entity_id)

        # Load original document
        doc_path = workspace.workspace_path / "documents" / "tender_original.txt"
        original_text = ""
        if doc_path.exists():
            original_text = doc_path.read_text(encoding='utf-8')

        # Get IC responses
        ic_responses = {}
        ic_state = self.workspace_manager.get_ic_completion_state(workspace_id)
        if ic_state:
            for field_text, qs in ic_state.get('question_status', {}).items():
                if qs.get('status') == 'completed' and qs.get('response'):
                    ic_responses[field_text] = qs['response']

        # Build completed document
        completed_text = self.build_completed_document(
            original_text, workspace.mentions, ic_responses, entity
        )

        # Build markdown output
        lines = []
        lines.append(f"# Proposal: {workspace.metadata.tender_name}")
        lines.append(f"")
        lines.append(f"**Workspace:** {workspace.metadata.workspace_name}")
        if workspace.metadata.tender_reference:
            lines.append(f"**RFT Reference:** {workspace.metadata.tender_reference}")
        if entity:
            lines.append(f"**Entity:** {entity.company.legal_name}")
        lines.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        lines.append(f"")
        lines.append("---")
        lines.append("")

        # Add completed document
        lines.append("## Completed Document")
        lines.append("")
        lines.append(completed_text)
        lines.append("")

        # Add Tier 2 responses section
        if ic_responses:
            lines.append("---")
            lines.append("")
            lines.append("## Substantive Responses")
            lines.append("")
            for field_text, response in ic_responses.items():
                lines.append(f"### {field_text}")
                lines.append("")
                lines.append(response)
                lines.append("")

                if include_citations and ic_state:
                    qs = ic_state.get('question_status', {}).get(field_text, {})
                    evidence = qs.get('evidence', [])
                    if evidence:
                        lines.append("**Sources:**")
                        for ev in evidence[:3]:
                            source = ev.get('source_doc', 'Unknown')
                            score = ev.get('relevance_score', 0)
                            lines.append(f"- {source} ({score:.0%} relevance)")
                        lines.append("")

        # Flag incomplete items
        if flag_incomplete:
            report = self.analyze_completeness(workspace_id)
            pending_items = [i for i in report.items if i.status == "pending"]
            if pending_items:
                lines.append("---")
                lines.append("")
                lines.append("## Incomplete Items")
                lines.append("")
                for item in pending_items:
                    tier_label = "Template Field" if item.tier == "tier1" else "AI Response"
                    lines.append(f"- [{tier_label}] {item.field_text}")
                lines.append("")

            if report.quality_flags:
                lines.append("## Quality Flags")
                lines.append("")
                for flag in report.quality_flags:
                    lines.append(f"- {flag}")
                lines.append("")

        return "\n".join(lines)

    def generate_export_docx(
        self,
        workspace_id: str,
        include_citations: bool = False,
        flag_incomplete: bool = True
    ) -> Optional[bytes]:
        """
        Generate a DOCX export of the proposal.

        Returns bytes of the DOCX file, or None if python-docx is not available.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
        except ImportError:
            logger.warning("python-docx not installed, DOCX export unavailable")
            return None

        workspace = self.workspace_manager.get_workspace(workspace_id)
        if not workspace:
            return None

        entity = None
        if workspace.metadata.entity_id:
            entity = self.entity_manager.get_entity_profile(workspace.metadata.entity_id)

        # Get IC responses
        ic_responses = {}
        ic_state = self.workspace_manager.get_ic_completion_state(workspace_id)
        if ic_state:
            for field_text, qs in ic_state.get('question_status', {}).items():
                if qs.get('status') == 'completed' and qs.get('response'):
                    ic_responses[field_text] = qs['response']

        # Load original document
        doc_path = workspace.workspace_path / "documents" / "tender_original.txt"
        original_text = ""
        if doc_path.exists():
            original_text = doc_path.read_text(encoding='utf-8')

        completed_text = self.build_completed_document(
            original_text, workspace.mentions, ic_responses, entity
        )

        # Build DOCX
        doc = Document()

        # Title
        title = doc.add_heading(f"Proposal: {workspace.metadata.tender_name}", level=0)

        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"Workspace: {workspace.metadata.workspace_name}\n").bold = True
        if workspace.metadata.tender_reference:
            meta_para.add_run(f"RFT Reference: {workspace.metadata.tender_reference}\n")
        if entity:
            meta_para.add_run(f"Entity: {entity.company.legal_name}\n")
        meta_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")

        doc.add_page_break()

        # Completed document content
        doc.add_heading("Completed Document", level=1)
        for paragraph_text in completed_text.split('\n\n'):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())

        # Substantive responses
        if ic_responses:
            doc.add_page_break()
            doc.add_heading("Substantive Responses", level=1)

            for field_text, response in ic_responses.items():
                doc.add_heading(field_text, level=2)
                for para_text in response.split('\n\n'):
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())

                if include_citations and ic_state:
                    qs = ic_state.get('question_status', {}).get(field_text, {})
                    evidence = qs.get('evidence', [])
                    if evidence:
                        sources_para = doc.add_paragraph()
                        sources_para.add_run("Sources: ").bold = True
                        for ev in evidence[:3]:
                            source = ev.get('source_doc', 'Unknown')
                            sources_para.add_run(f"{source}; ")

        # Incomplete items
        if flag_incomplete:
            report = self.analyze_completeness(workspace_id)
            pending_items = [i for i in report.items if i.status == "pending"]
            if pending_items:
                doc.add_page_break()
                doc.add_heading("Incomplete Items", level=1)
                for item in pending_items:
                    tier_label = "Template Field" if item.tier == "tier1" else "AI Response"
                    doc.add_paragraph(f"[{tier_label}] {item.field_text}", style='List Bullet')

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
