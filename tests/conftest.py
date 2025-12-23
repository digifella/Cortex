"""
Pytest Configuration and Shared Fixtures
Version: 1.0.0
Purpose: Provide reusable test fixtures and configuration for all tests
"""

import pytest
import tempfile
import shutil
from pathlib import Path
from typing import Generator
import sys

# Add project root to path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from cortex_engine.config import get_cortex_config
from cortex_engine.version_config import CORTEX_VERSION


# ============================================================================
# Session-scoped fixtures (run once per test session)
# ============================================================================

@pytest.fixture(scope="session")
def project_root_path() -> Path:
    """Return the project root directory path."""
    return Path(__file__).parent.parent


@pytest.fixture(scope="session")
def cortex_version() -> str:
    """Return the current Cortex version."""
    return CORTEX_VERSION


@pytest.fixture(scope="session")
def cortex_config() -> dict:
    """Return the Cortex configuration dictionary."""
    return get_cortex_config()


# ============================================================================
# Function-scoped fixtures (run for each test)
# ============================================================================

@pytest.fixture
def temp_dir() -> Generator[Path, None, None]:
    """Create a temporary directory that is cleaned up after the test."""
    temp_path = Path(tempfile.mkdtemp(prefix="cortex_test_"))
    yield temp_path
    # Cleanup
    if temp_path.exists():
        shutil.rmtree(temp_path)


@pytest.fixture
def temp_db_path(temp_dir: Path) -> Path:
    """Create a temporary database path for testing."""
    db_path = temp_dir / "test_knowledge_hub_db"
    db_path.mkdir(parents=True, exist_ok=True)
    return db_path


@pytest.fixture
def temp_source_path(temp_dir: Path) -> Path:
    """Create a temporary source documents path for testing."""
    source_path = temp_dir / "test_sources"
    source_path.mkdir(parents=True, exist_ok=True)
    return source_path


@pytest.fixture
def sample_text_file(temp_source_path: Path) -> Path:
    """Create a sample text file for testing."""
    file_path = temp_source_path / "sample.txt"
    file_path.write_text("This is a sample text file for testing purposes.")
    return file_path


@pytest.fixture
def sample_docx_file(temp_source_path: Path) -> Path:
    """Create a sample DOCX file for testing."""
    try:
        import docx
        doc = docx.Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test paragraph.")
        doc.add_paragraph("This is another test paragraph.")

        file_path = temp_source_path / "sample.docx"
        doc.save(str(file_path))
        return file_path
    except ImportError:
        pytest.skip("python-docx not available")


@pytest.fixture
def sample_pdf_file(temp_source_path: Path) -> Path:
    """Create a sample PDF file for testing."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "This is a test PDF document.")

        file_path = temp_source_path / "sample.pdf"
        doc.save(str(file_path))
        doc.close()
        return file_path
    except ImportError:
        pytest.skip("PyMuPDF not available")


# ============================================================================
# Mock fixtures for external services
# ============================================================================

@pytest.fixture
def mock_ollama_client(monkeypatch):
    """Mock Ollama client for testing without requiring Ollama service."""
    class MockOllamaClient:
        def list(self):
            return {
                'models': [
                    {'name': 'mistral:latest'},
                    {'name': 'llava:7b'},
                ]
            }

        def generate(self, model, prompt, **kwargs):
            return {
                'response': 'This is a mock response from Ollama'
            }

    import ollama
    monkeypatch.setattr(ollama, "Client", MockOllamaClient)
    return MockOllamaClient()


@pytest.fixture
def mock_chromadb_client(monkeypatch):
    """Mock ChromaDB client for testing without requiring ChromaDB."""
    class MockCollection:
        def __init__(self):
            self._documents = []

        def add(self, ids, documents, metadatas=None, embeddings=None):
            for i, doc_id in enumerate(ids):
                self._documents.append({
                    'id': doc_id,
                    'document': documents[i] if documents else None,
                    'metadata': metadatas[i] if metadatas else {},
                })

        def query(self, query_texts, n_results=5, where=None, **kwargs):
            return {
                'ids': [['doc1', 'doc2']],
                'documents': [['Sample doc 1', 'Sample doc 2']],
                'metadatas': [[{}, {}]],
                'distances': [[0.1, 0.2]],
            }

        def count(self):
            return len(self._documents)

        def get(self, ids=None, where=None, limit=None, **kwargs):
            if ids:
                docs = [d for d in self._documents if d['id'] in ids]
            else:
                docs = self._documents

            return {
                'ids': [d['id'] for d in docs],
                'documents': [d['document'] for d in docs],
                'metadatas': [d['metadata'] for d in docs],
            }

        def delete(self, ids=None, where=None):
            if ids:
                self._documents = [d for d in self._documents if d['id'] not in ids]

    class MockChromaClient:
        def __init__(self, *args, **kwargs):
            self._collections = {}

        def get_or_create_collection(self, name, **kwargs):
            if name not in self._collections:
                self._collections[name] = MockCollection()
            return self._collections[name]

        def list_collections(self):
            return [{'name': name} for name in self._collections.keys()]

        def delete_collection(self, name):
            if name in self._collections:
                del self._collections[name]

    import chromadb
    monkeypatch.setattr(chromadb, "PersistentClient", MockChromaClient)
    return MockChromaClient()


# ============================================================================
# Test environment setup
# ============================================================================

@pytest.fixture(autouse=True)
def reset_environment():
    """Reset environment variables before each test."""
    import os
    original_env = os.environ.copy()
    yield
    # Restore original environment
    os.environ.clear()
    os.environ.update(original_env)


# ============================================================================
# Pytest hooks for custom behavior
# ============================================================================

def pytest_configure(config):
    """Configure pytest with custom settings."""
    config.addinivalue_line(
        "markers", "smoke: Quick smoke tests to verify basic functionality"
    )


def pytest_collection_modifyitems(config, items):
    """Modify test collection to add markers automatically."""
    for item in items:
        # Auto-mark slow tests
        if hasattr(item, 'callspec') and 'slow' in str(item.callspec):
            item.add_marker(pytest.mark.slow)

        # Auto-mark tests in integration directory
        if "integration" in str(item.fspath):
            item.add_marker(pytest.mark.integration)

        # Auto-mark tests in ui directory
        if "ui" in str(item.fspath):
            item.add_marker(pytest.mark.ui)


def pytest_report_header(config):
    """Add custom header to test report."""
    return [
        f"Cortex Suite Test Suite v{CORTEX_VERSION}",
        f"Project root: {project_root}",
    ]
