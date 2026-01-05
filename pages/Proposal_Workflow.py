"""
Proposal Workflow - Complete Tender Response System
Version: 2.0.0
Date: 2026-01-03

Purpose: Complete workflow for responding to tenders with workspace-based extraction.

Workflow:
1. Select tender document (file picker) → Create workspace
2. Select source documents for extraction (file picker)
3. Extract structured data from sources → Populate workspace collection
4. Match tender fields to extracted data (Phase 2)
5. Review and approve matches
6. Fill and export completed tender (Phase 3)

New in v2.0.0:
- Workspace model: Each tender gets its own ChromaDB collection
- Progressive additions: Add research, notes, narratives as you work
- Hybrid data: Structured JSON + unstructured workspace collection
"""

import streamlit as st
import sys
from pathlib import Path
from datetime import datetime
import asyncio
import tempfile
import shutil

# Set page config
st.set_page_config(
    page_title="Proposal Workflow",
    page_icon="📝",
    layout="wide"
)

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from cortex_engine.entity_manager import EntityManager, EntityProfile, EntityType
from cortex_engine.workspace_manager import WorkspaceManager
from cortex_engine.workspace_schema import WorkspaceStatus, DocumentSource
from cortex_engine.tender_data_extractor import TenderDataExtractor
from cortex_engine.adaptive_model_manager import AdaptiveModelManager
from cortex_engine.config_manager import ConfigManager
from cortex_engine.collection_manager import WorkingCollectionManager
from cortex_engine.ui_theme import apply_theme, section_header
from cortex_engine.utils import (
    convert_windows_to_wsl_path,
    get_logger
)
# Phase 2: Field Matching imports
from cortex_engine.proposals.tender_field_parser import TenderFieldParser
from cortex_engine.proposals.field_matcher import FieldMatcher
from cortex_engine.proposals.field_classifier import get_classifier
import chromadb
from chromadb.config import Settings as ChromaSettings
import networkx as nx
import docx

# Apply theme
apply_theme()

# Set up logging
logger = get_logger(__name__)

# Page version
PAGE_VERSION = "v2.0.0"


def initialize_session_state():
    """Initialize session state variables."""
    if 'workspace_manager' not in st.session_state:
        st.session_state.workspace_manager = None

    if 'current_workspace' not in st.session_state:
        st.session_state.current_workspace = None

    if 'tender_document' not in st.session_state:
        st.session_state.tender_document = None

    if 'tender_filename' not in st.session_state:
        st.session_state.tender_filename = None

    if 'tender_id' not in st.session_state:
        st.session_state.tender_id = None

    if 'source_documents' not in st.session_state:
        st.session_state.source_documents = []

    if 'extracted_data' not in st.session_state:
        st.session_state.extracted_data = None

    if 'selected_entity' not in st.session_state:
        st.session_state.selected_entity = None

    # Step 4: Field Matching state
    if 'field_mappings' not in st.session_state:
        st.session_state.field_mappings = None

    if 'detected_fields' not in st.session_state:
        st.session_state.detected_fields = None

    if 'matching_in_progress' not in st.session_state:
        st.session_state.matching_in_progress = False

    if 'matching_log_messages' not in st.session_state:
        st.session_state.matching_log_messages = []

    if 'filter_status' not in st.session_state:
        st.session_state.filter_status = "All"

    if 'filter_confidence' not in st.session_state:
        st.session_state.filter_confidence = 0.0

    if 'workflow_step' not in st.session_state:
        st.session_state.workflow_step = 1


def render_workflow_progress(current_step: int):
    """Render workflow progress indicator."""
    steps = [
        "1️⃣ Select Tender",
        "2️⃣ Select Sources",
        "3️⃣ Extract Data",
        "4️⃣ Match Fields",
        "5️⃣ Fill & Export"
    ]

    cols = st.columns(len(steps))
    for i, (col, step_name) in enumerate(zip(cols, steps), 1):
        with col:
            if i < current_step:
                st.success(f"✅ {step_name}")
            elif i == current_step:
                st.info(f"▶️ {step_name}")
            else:
                st.caption(f"⚪ {step_name}")


def render_step1_select_tender():
    """Step 1: Select tender document to fill out."""
    section_header("📄", "Step 1: Select Tender Document", "Choose the RFT/RFQ you need to complete")

    st.markdown("""
    Upload the tender document you need to respond to (e.g., RFT12493.docx).
    This is the document with blank fields that need to be filled.
    """)

    # Option 1: Upload tender file
    st.markdown("### 📤 Upload Tender Document")
    uploaded_tender = st.file_uploader(
        "Choose tender document (.docx)",
        type=['docx'],
        key="tender_uploader",
        help="Select the tender/RFQ/RFT document you need to complete"
    )

    if uploaded_tender:
        # Save to session state
        st.session_state.tender_filename = uploaded_tender.name

        # Extract tender ID from filename (e.g., "RFT12493" from "RFT12493-Request-for-Tender.docx")
        if not st.session_state.tender_id:
            # Try to extract tender ID from filename
            import re
            match = re.search(r'(RFT|RFQ|RFP|EOI)[\s-]?(\d+)', uploaded_tender.name, re.IGNORECASE)
            if match:
                st.session_state.tender_id = f"{match.group(1).upper()}{match.group(2)}"
            else:
                # Fallback: use filename without extension
                st.session_state.tender_id = Path(uploaded_tender.name).stem[:20]

        # Load document
        try:
            tender_doc = docx.Document(uploaded_tender)
            st.session_state.tender_document = tender_doc

            # Show document info
            st.success(f"✅ Loaded: {uploaded_tender.name}")
            st.info(f"🆔 Tender ID: **{st.session_state.tender_id}**")

            # Count sections
            para_count = len(tender_doc.paragraphs)
            table_count = len(tender_doc.tables)

            col1, col2 = st.columns(2)
            with col1:
                st.metric("Paragraphs", para_count)
            with col2:
                st.metric("Tables", table_count)

            # Preview
            with st.expander("📖 Preview First Few Paragraphs"):
                for i, para in enumerate(tender_doc.paragraphs[:10]):
                    if para.text.strip():
                        st.caption(f"{i+1}. {para.text[:200]}...")

            # Next step button - creates workspace
            if st.button("➡️ Create Workspace & Continue", type="primary"):
                # Create workspace if not already created
                if not st.session_state.current_workspace:
                    try:
                        workspace = st.session_state.workspace_manager.create_workspace(
                            tender_id=st.session_state.tender_id,
                            tender_filename=st.session_state.tender_filename
                        )
                        st.session_state.current_workspace = workspace
                        st.success(f"✅ Created workspace: {workspace.workspace_id}")
                        logger.info(f"Created workspace {workspace.workspace_id}")
                    except Exception as e:
                        st.error(f"Failed to create workspace: {str(e)}")
                        logger.error(f"Workspace creation failed: {e}", exc_info=True)
                        return

                st.session_state.workflow_step = 2
                st.rerun()

        except Exception as e:
            st.error(f"❌ Failed to load document: {str(e)}")
            logger.error(f"Tender document load failed: {e}", exc_info=True)

    # Option 2: Use entity's saved tender template
    st.divider()
    st.markdown("### 🏢 Or Use Entity Template")
    st.info("💡 **Coming Soon:** Select from saved tender templates in your entity profiles")


def render_step2_select_sources():
    """Step 2: Select source documents for data extraction."""
    section_header("📁", "Step 2: Select Source Documents", "Choose documents to extract company data from")

    st.markdown(f"""
    **Selected Tender:** {st.session_state.tender_filename}

    Now select the source documents containing your company information:
    - Company registration (ABN, ACN)
    - Insurance certificates
    - Team CVs/qualifications
    - Project case studies
    - References
    """)

    # Option 1: Upload source files
    st.markdown("### 📤 Upload Source Documents")
    uploaded_sources = st.file_uploader(
        "Choose source documents",
        type=['pdf', 'docx', 'txt'],
        accept_multiple_files=True,
        key="source_uploader",
        help="Select documents containing company information to extract"
    )

    if uploaded_sources:
        st.success(f"✅ Selected {len(uploaded_sources)} source documents")

        # Show list
        with st.expander("📋 Source Documents"):
            for doc in uploaded_sources:
                st.caption(f"📄 {doc.name} ({doc.size} bytes)")

        st.session_state.source_documents = uploaded_sources

    # Option 2: Select from Working Collections
    st.divider()
    st.markdown("### 📚 Or Select from Working Collections")

    try:
        # Cache collection loading for performance
        if 'working_collection_mgr' not in st.session_state:
            with st.spinner("Loading collections..."):
                st.session_state.working_collection_mgr = WorkingCollectionManager()

        collection_mgr = st.session_state.working_collection_mgr
        working_collections = list(collection_mgr.collections.values())

        if working_collections:
            # Initialize selected collections in session state
            if 'selected_workflow_collections' not in st.session_state:
                st.session_state.selected_workflow_collections = []

            st.markdown("**Available Collections:**")

            # Sort collections by name
            sorted_collections = sorted(working_collections, key=lambda c: c.get('name', '').lower())

            for collection in sorted_collections:
                collection_name = collection.get('name', '')
                doc_ids = collection.get('doc_ids', [])

                col_check, col_name, col_count = st.columns([1, 6, 2])

                with col_check:
                    is_selected = collection_name in st.session_state.selected_workflow_collections
                    if st.checkbox(
                        "",
                        value=is_selected,
                        key=f"workflow_coll_{collection_name}",
                        label_visibility="collapsed"
                    ):
                        if collection_name not in st.session_state.selected_workflow_collections:
                            st.session_state.selected_workflow_collections.append(collection_name)
                    else:
                        if collection_name in st.session_state.selected_workflow_collections:
                            st.session_state.selected_workflow_collections.remove(collection_name)

                with col_name:
                    st.markdown(f"📚 **{collection_name}**")

                with col_count:
                    st.caption(f"{len(doc_ids)} docs")

            # Show selected summary
            if st.session_state.selected_workflow_collections:
                total_docs = sum(
                    len(collection_mgr.collections[name].get('doc_ids', []))
                    for name in st.session_state.selected_workflow_collections
                    if name in collection_mgr.collections
                )
                st.success(f"✅ Selected {len(st.session_state.selected_workflow_collections)} collection(s) with {total_docs} documents")

                # Store collection selections in workflow state
                st.session_state.source_collections = st.session_state.selected_workflow_collections
                st.session_state.source_collection_doc_ids = []
                for coll_name in st.session_state.selected_workflow_collections:
                    coll_data = collection_mgr.collections.get(coll_name)
                    if coll_data:
                        st.session_state.source_collection_doc_ids.extend(coll_data.get('doc_ids', []))
        else:
            st.info("No working collections found. Create collections in the Collection Management page.")

    except Exception as e:
        st.error(f"Failed to load working collections: {e}")
        logger.error(f"Working collection loading failed in workflow: {e}", exc_info=True)

    # Option 3: Browse KB Folders
    st.divider()
    st.markdown("### 📁 Or Browse KB Folders")

    # Get config
    config_manager = ConfigManager()
    config = config_manager.get_config()
    db_path = config.get('ai_database_path')

    if db_path:
        try:
            wsl_db_path = convert_windows_to_wsl_path(db_path)

            # Initialize KB navigator if not already done
            if 'kb_navigator' not in st.session_state or st.session_state.kb_navigator is None:
                from cortex_engine.kb_navigator import KBNavigator
                st.session_state.kb_navigator = KBNavigator(
                    db_path=wsl_db_path,
                    collection_name=config.get('collection_name', 'knowledge_hub_collection')
                )

            kb_navigator = st.session_state.kb_navigator

            # Initialize session state for folder expansion
            if 'expanded_folders' not in st.session_state:
                st.session_state.expanded_folders = set()

            if 'selected_kb_folders' not in st.session_state:
                st.session_state.selected_kb_folders = []

            # Build folder tree
            root_folder = kb_navigator.build_folder_tree()

            def render_folder_tree(folder_node, level=0, parent_path=""):
                """Recursively render folder tree with expand/collapse."""
                indent = "　" * level  # Japanese space for indentation
                folder_path = folder_node.full_path if folder_node.full_path else "/"

                # Calculate total documents including subfolders
                total_docs = folder_node.get_total_document_count()

                # Create unique key for this folder
                folder_key = f"folder_{folder_path}_{level}"

                # Column layout: expand button, checkbox, folder name
                col_expand, col_check, col_name = st.columns([1, 1, 8])

                with col_expand:
                    if folder_node.subfolders:
                        # Has subfolders - show expand/collapse button
                        is_expanded = folder_path in st.session_state.expanded_folders
                        expand_label = "▼" if is_expanded else "▶"
                        if st.button(expand_label, key=f"expand_{folder_key}"):
                            if is_expanded:
                                st.session_state.expanded_folders.remove(folder_path)
                            else:
                                st.session_state.expanded_folders.add(folder_path)
                            st.rerun()
                    else:
                        st.write("　")  # Empty space if no subfolders

                with col_check:
                    # Checkbox to select this folder
                    is_selected = folder_path in st.session_state.selected_kb_folders
                    if st.checkbox(
                        "",
                        value=is_selected,
                        key=f"select_{folder_key}",
                        label_visibility="collapsed"
                    ):
                        if folder_path not in st.session_state.selected_kb_folders:
                            st.session_state.selected_kb_folders.append(folder_path)
                    else:
                        if folder_path in st.session_state.selected_kb_folders:
                            st.session_state.selected_kb_folders.remove(folder_path)

                with col_name:
                    # Folder name and document count
                    folder_display = folder_node.name if folder_node.name else "Root"
                    st.markdown(f"{indent}📁 **{folder_display}** ({total_docs} docs)")

                # Render subfolders if expanded
                if folder_path in st.session_state.expanded_folders and folder_node.subfolders:
                    for subfolder_name in sorted(folder_node.subfolders.keys()):
                        subfolder = folder_node.subfolders[subfolder_name]
                        render_folder_tree(subfolder, level + 1, folder_path)

            # Render the folder tree
            st.markdown("**📂 Knowledge Base Folders:**")
            with st.container():
                render_folder_tree(root_folder)

            # Show selected folders summary
            if st.session_state.selected_kb_folders:
                st.divider()
                st.success(f"✅ Selected {len(st.session_state.selected_kb_folders)} folder(s)")

                # Calculate total documents
                total_docs = 0
                all_doc_ids = []
                for folder_path in st.session_state.selected_kb_folders:
                    docs = kb_navigator.filter_by_folder(folder_path, include_subfolders=True)
                    total_docs += len(docs)
                    all_doc_ids.extend([doc.doc_id for doc in docs])

                st.info(f"📊 Total documents across all selected folders: {total_docs}")

                # Show selected folder list
                with st.expander("📋 Selected Folders"):
                    for folder_path in st.session_state.selected_kb_folders:
                        folder_node = kb_navigator.get_folder_node(folder_path)
                        doc_count = folder_node.get_total_document_count() if folder_node else 0
                        st.caption(f"📁 {folder_path} ({doc_count} docs)")

                # Save document IDs to session state
                st.session_state.selected_kb_doc_ids = all_doc_ids

        except Exception as e:
            logger.error(f"KB folder browsing failed: {e}", exc_info=True)
            st.error(f"Failed to load KB folders: {str(e)}")

    # Option 3: Use pre-configured entity
    st.divider()
    st.markdown("### 🏢 Or Use Entity Profile")

    # Get config
    config_manager = ConfigManager()
    config = config_manager.get_config()
    db_path = config.get('ai_database_path')

    if db_path:
        try:
            wsl_db_path = convert_windows_to_wsl_path(db_path)
            entity_manager = EntityManager(Path(wsl_db_path))
            entities = entity_manager.list_entities()

            if entities:
                entity_names = [e.entity_name for e in entities]
                selected_entity_name = st.selectbox(
                    "Select entity with pre-configured sources",
                    options=["None"] + entity_names,
                    help="Use documents from a pre-configured entity profile"
                )

                if selected_entity_name != "None":
                    entity = next(e for e in entities if e.entity_name == selected_entity_name)
                    st.session_state.selected_entity = entity

                    # Update workspace with entity info
                    if st.session_state.current_workspace:
                        workspace = st.session_state.current_workspace
                        if not workspace.entity_id or workspace.entity_id != entity.entity_id:
                            # Update workspace metadata with entity
                            workspace.entity_id = entity.entity_id
                            workspace.entity_name = entity.entity_name
                            workspace.workspace_name = f"{workspace.tender_id} - {entity.entity_name}"
                            st.session_state.workspace_manager._save_metadata(
                                workspace.workspace_id,
                                workspace
                            )
                            st.success(f"✅ Linked workspace to entity: {entity.entity_name}")

                    st.info(f"""
                    ✅ Using **{entity.entity_name}**
                    - Source folders: {len(entity.source_folders)}
                    - Source documents: {entity.source_document_count}
                    - Last extracted: {entity.last_extracted.strftime('%Y-%m-%d') if entity.last_extracted else 'Never'}
                    """)
            else:
                st.warning("No entities configured. Go to **Proposal Entity Manager** to create one.")

        except Exception as e:
            logger.warning(f"Could not load entities: {e}")
            st.warning("Could not load entity profiles")

    # Navigation
    st.divider()
    col1, col2 = st.columns(2)

    with col1:
        if st.button("⬅️ Back to Tender Selection"):
            st.session_state.workflow_step = 1
            st.rerun()

    with col2:
        has_kb_folders = (
            'selected_kb_folders' in st.session_state and
            len(st.session_state.selected_kb_folders) > 0
        )

        has_working_collections = (
            'source_collection_doc_ids' in st.session_state and
            len(st.session_state.source_collection_doc_ids) > 0
        )

        can_continue = (
            (uploaded_sources and len(uploaded_sources) > 0) or
            st.session_state.selected_entity is not None or
            has_kb_folders or
            has_working_collections
        )

        if st.button("➡️ Continue to Extract Data", type="primary", disabled=not can_continue):
            st.session_state.workflow_step = 3
            st.rerun()

        if not can_continue:
            st.caption("⚠️ Please select source documents, working collections, KB folders, or choose an entity")


def render_step3_extract_data():
    """Step 3: Extract structured data from source documents."""
    section_header("🔍", "Step 3: Extract Structured Data", "Extract company information from sources")

    # Determine source description
    if st.session_state.selected_entity:
        source_desc = f"Entity: {st.session_state.selected_entity.entity_name} ({st.session_state.selected_entity.source_document_count} docs)"
    elif 'source_collection_doc_ids' in st.session_state and st.session_state.source_collection_doc_ids:
        total_docs = len(st.session_state.source_collection_doc_ids)
        num_collections = len(st.session_state.get('source_collections', []))
        source_desc = f"{num_collections} working collection(s) ({total_docs} docs)"
    elif 'selected_kb_folders' in st.session_state and st.session_state.selected_kb_folders:
        total_docs = len(st.session_state.selected_kb_doc_ids) if 'selected_kb_doc_ids' in st.session_state else 0
        source_desc = f"{len(st.session_state.selected_kb_folders)} KB folders ({total_docs} docs)"
    elif st.session_state.source_documents:
        source_desc = f"{len(st.session_state.source_documents)} uploaded files"
    else:
        source_desc = "None selected"

    st.markdown(f"""
    **Selected Tender:** {st.session_state.tender_filename}

    **Sources:** {source_desc}
    """)

    # Show what will be extracted
    st.markdown("### 📊 Extraction Target")
    st.info("""
    The system will extract:
    - 🏢 Organization details (ABN, ACN, address, contact)
    - 🛡️ Insurance policies (policy numbers, coverage, expiry)
    - 🎓 Team qualifications and certifications
    - 💼 Work experience and employment history
    - 🚀 Project experience and case studies
    - 📞 Client references
    - ⭐ Organizational capabilities
    """)

    # Extract button
    if st.button("🚀 Extract Structured Data", type="primary", key="extract_data_button"):
        with st.spinner("🔍 Extracting structured data..."):
            try:
                # Get config
                config_manager = ConfigManager()
                config = config_manager.get_config()
                db_path = config.get('ai_database_path')
                wsl_db_path = convert_windows_to_wsl_path(db_path)

                # Initialize extractor
                chroma_db_path = str(Path(wsl_db_path) / "knowledge_hub_db")
                db_settings = ChromaSettings(anonymized_telemetry=False)
                chroma_client = chromadb.PersistentClient(path=chroma_db_path, settings=db_settings)
                collection = chroma_client.get_collection(config.get('collection_name', 'knowledge_hub'))

                # Load knowledge graph
                graph_path = Path(wsl_db_path) / "knowledge_cortex.gpickle"
                knowledge_graph = None
                if graph_path.exists():
                    try:
                        import pickle
                        with open(graph_path, 'rb') as f:
                            knowledge_graph = pickle.load(f)
                    except Exception as e:
                        logger.warning(f"Could not load knowledge graph: {e}")
                        knowledge_graph = None

                model_manager = AdaptiveModelManager()
                extractor = TenderDataExtractor(
                    vector_index=collection,
                    knowledge_graph=knowledge_graph,
                    model_manager=model_manager,
                    db_path=Path(wsl_db_path)
                )

                # Progress tracking
                progress_placeholder = st.empty()

                def progress_callback(message):
                    progress_placeholder.info(message)

                # Run extraction
                if st.session_state.selected_entity:
                    # Use entity's pre-configured documents
                    entity = st.session_state.selected_entity
                    structured_data = asyncio.run(
                        extractor.extract_all_structured_data(
                            progress_callback=progress_callback,
                            entity_id=entity.entity_id,
                            document_ids=entity.source_document_ids
                        )
                    )
                elif 'source_collection_doc_ids' in st.session_state and st.session_state.source_collection_doc_ids:
                    # Use selected working collections
                    num_collections = len(st.session_state.get('source_collections', []))
                    collection_names = ", ".join(st.session_state.get('source_collections', []))
                    progress_callback(f"📚 Extracting from {num_collections} working collection(s): {collection_names}...")

                    # Remove duplicates while preserving order
                    unique_doc_ids = list(dict.fromkeys(st.session_state.source_collection_doc_ids))

                    structured_data = asyncio.run(
                        extractor.extract_all_structured_data(
                            progress_callback=progress_callback,
                            document_ids=unique_doc_ids
                        )
                    )
                elif 'selected_kb_doc_ids' in st.session_state and st.session_state.selected_kb_doc_ids:
                    # Use selected KB folders
                    progress_callback(f"📁 Extracting from {len(st.session_state.selected_kb_folders)} KB folder(s)...")
                    structured_data = asyncio.run(
                        extractor.extract_all_structured_data(
                            progress_callback=progress_callback,
                            document_ids=st.session_state.selected_kb_doc_ids
                        )
                    )
                else:
                    # Use uploaded source documents
                    # TODO: Ingest uploaded documents to temp collection first
                    st.warning("⚠️ Direct file upload extraction not yet implemented. Please use KB folders or entity profiles for now.")
                    return

                # Save to session state
                st.session_state.extracted_data = structured_data

                # Populate workspace with extracted data
                if st.session_state.current_workspace:
                    progress_placeholder.info("📦 Populating workspace with extracted data...")
                    workspace_manager = st.session_state.workspace_manager
                    workspace_id = st.session_state.current_workspace.workspace_id

                    success = asyncio.run(
                        extractor.populate_workspace_with_extraction(
                            workspace_manager=workspace_manager,
                            workspace_id=workspace_id,
                            structured_data=structured_data
                        )
                    )

                    if success:
                        # Update workspace status
                        workspace_manager.update_workspace_status(
                            workspace_id,
                            WorkspaceStatus.IN_PROGRESS
                        )
                        # Refresh workspace metadata
                        st.session_state.current_workspace = workspace_manager.get_workspace(workspace_id)
                        progress_placeholder.success("✅ Workspace populated with extracted data!")
                    else:
                        progress_placeholder.warning("⚠️ Extraction succeeded but workspace population failed")

                # Show results
                st.success("✅ Extraction complete!")

                stats = structured_data.summary_stats
                st.markdown("### 📊 Extraction Summary")
                col1, col2, col3, col4 = st.columns(4)

                with col1:
                    st.metric("Organization", "✓" if structured_data.organization else "✗")
                    st.metric("Insurances", stats.get('insurances', 0))

                with col2:
                    st.metric("Qualifications", stats.get('qualifications', 0))
                    st.metric("Work Experience", stats.get('work_experiences', 0))

                with col3:
                    st.metric("Projects", stats.get('projects', 0))
                    st.metric("References", stats.get('references', 0))

                with col4:
                    st.metric("Capabilities", stats.get('capabilities', 0))

                # Show workspace info
                if st.session_state.current_workspace:
                    workspace = st.session_state.current_workspace
                    st.divider()
                    st.markdown("### 📁 Workspace Information")
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Workspace ID", workspace.workspace_id)
                    with col2:
                        st.metric("Documents in Workspace", workspace.document_count)
                    with col3:
                        st.metric("Status", workspace.status.replace('_', ' ').title())

                # Next step
                st.info("✨ Ready to match tender fields! Click 'Continue' below.")

            except Exception as e:
                st.error(f"❌ Extraction failed: {str(e)}")
                logger.error(f"Extraction failed: {e}", exc_info=True)

    # Show extracted data if available
    if st.session_state.extracted_data:
        with st.expander("📋 View Extracted Data"):
            data = st.session_state.extracted_data

            if data.organization:
                st.markdown("**Organization:**")
                st.json({
                    "legal_name": data.organization.legal_name,
                    "abn": data.organization.abn,
                    "address": data.organization.address
                })

            if data.insurances:
                st.markdown(f"**Insurances ({len(data.insurances)}):**")
                for ins in data.insurances[:3]:
                    st.caption(f"• {ins.insurance_type.value}: {ins.policy_number} (Expires: {ins.expiry_date})")

    # Navigation
    st.divider()
    col1, col2 = st.columns(2)

    with col1:
        if st.button("⬅️ Back to Source Selection"):
            st.session_state.workflow_step = 2
            st.rerun()

    with col2:
        can_continue = st.session_state.extracted_data is not None

        if st.button("➡️ Continue to Match Fields", type="primary", disabled=not can_continue):
            st.session_state.workflow_step = 4
            st.rerun()

        if not can_continue:
            st.caption("⚠️ Please extract data first")


def render_step4_match_fields():
    """Step 4: Match tender fields to extracted data."""
    section_header("🎯", "Step 4: Match Tender Fields", "Auto-match tender fields to extracted data")

    # Check if we have required data
    if not st.session_state.current_workspace:
        st.error("No workspace found. Please complete Steps 1-3 first.")
        if st.button("⬅️ Back to Step 1"):
            st.session_state.workflow_step = 1
            st.rerun()
        return

    if not st.session_state.tender_document:
        st.error("No tender document found. Please upload a tender in Step 1.")
        if st.button("⬅️ Back to Step 1"):
            st.session_state.workflow_step = 1
            st.rerun()
        return

    if not st.session_state.extracted_data:
        st.error("No extracted data found. Please complete extraction in Step 3.")
        if st.button("⬅️ Back to Step 3"):
            st.session_state.workflow_step = 3
            st.rerun()
        return

    # Load or initialize field mappings
    workspace_id = st.session_state.current_workspace.workspace_id
    workspace_manager = st.session_state.workspace_manager

    if st.session_state.field_mappings is None:
        # Try to load existing mappings
        existing_mappings = workspace_manager.get_field_mappings(workspace_id)
        if existing_mappings:
            st.session_state.field_mappings = existing_mappings
            st.success(f"Loaded {len(existing_mappings)} existing field mappings")

    # Initial state: Show auto-match button
    if st.session_state.field_mappings is None:
        st.info("**Ready to analyze tender document and match fields**")

        st.markdown("""
        This step will:
        1. 🔍 **Parse** the tender document to detect fillable fields
        2. 🏷️ **Classify** field types (ABN, insurance, experience, etc.)
        3. 🎯 **Match** fields to your extracted data using AI
        4. 📊 **Score** matches with confidence levels
        5. ✅ **Review** - you approve, edit, or reject each match
        """)

        st.divider()

        col1, col2, col3 = st.columns([2, 1, 2])
        with col2:
            if st.button("🚀 Start Auto-Matching", type="primary", use_container_width=True):
                run_field_matching()

        st.divider()
        if st.button("⬅️ Back to Extraction"):
            st.session_state.workflow_step = 3
            st.rerun()

        return

    # We have field mappings - show the review interface
    render_matching_interface()


def add_matching_log(message: str):
    """Add message to matching log."""
    from datetime import datetime
    timestamp = datetime.now().strftime("%H:%M:%S")
    log_entry = f"[{timestamp}] {message}"
    st.session_state.matching_log_messages.append(log_entry)
    logger.info(message)


def run_field_matching():
    """Execute the field matching process with real-time logging."""
    workspace_id = st.session_state.current_workspace.workspace_id
    workspace_manager = st.session_state.workspace_manager

    # Clear previous logs
    st.session_state.matching_log_messages = []
    st.session_state.matching_in_progress = True

    # Create progress containers
    progress_container = st.empty()
    log_container_wrapper = st.container()

    try:
        with log_container_wrapper:
            with st.expander("📋 Processing Log", expanded=True):
                log_display = st.container(height=400, border=True)

        # Step 1: Parse tender document
        add_matching_log("🔍 Starting tender document parsing...")
        progress_container.info("🔍 Parsing tender document...")

        tender_doc_obj = st.session_state.tender_document
        parser = TenderFieldParser(use_classification=True)

        add_matching_log(f"  → Parser initialized with classification enabled")

        detected_fields = parser.parse_tender_document(tender_doc_obj)
        st.session_state.detected_fields = detected_fields

        add_matching_log(f"✅ Detected {len(detected_fields)} fillable fields")
        add_matching_log(f"  → Field types: Tables, Paragraphs, Placeholders")

        # Update log display
        with log_display:
            st.code("\n".join(st.session_state.matching_log_messages), language="log")

        # Step 2: Load entity data
        add_matching_log("📊 Loading entity data from workspace...")
        progress_container.info("📊 Loading entity data...")

        entity_data = workspace_manager.get_entity_snapshot(workspace_id)
        if not entity_data:
            add_matching_log("❌ ERROR: No entity data found in workspace")
            with log_display:
                st.code("\n".join(st.session_state.matching_log_messages), language="log")
            progress_container.error("No entity data found in workspace")
            st.session_state.matching_in_progress = False
            return

        # Log entity data summary
        org_name = entity_data.get('organization', {}).get('legal_name', 'Unknown')
        add_matching_log(f"✅ Loaded entity data for: {org_name}")

        if 'insurances' in entity_data:
            add_matching_log(f"  → Found {len(entity_data['insurances'])} insurance policies")
        if 'projects' in entity_data:
            add_matching_log(f"  → Found {len(entity_data['projects'])} project references")
        if 'capabilities' in entity_data:
            add_matching_log(f"  → Found {len(entity_data['capabilities'])} capabilities")

        with log_display:
            st.code("\n".join(st.session_state.matching_log_messages), language="log")

        # Step 3: Match fields
        add_matching_log("🎯 Starting intelligent field matching...")
        add_matching_log("  → Using hybrid approach: Structured + Semantic + LLM")
        progress_container.info("🎯 Matching fields to data (this may take a few minutes)...")

        with log_display:
            st.code("\n".join(st.session_state.matching_log_messages), language="log")

        matcher = FieldMatcher(
            entity_data=entity_data,
            workspace_id=workspace_id,
            model_name="qwen2.5:14b-instruct-q4_K_M"
        )

        add_matching_log(f"  → Model: qwen2.5:14b-instruct-q4_K_M")

        with log_display:
            st.code("\n".join(st.session_state.matching_log_messages), language="log")

        field_mappings = matcher.match_all_fields(detected_fields)

        # Summarize matches
        matched_count = sum(1 for fm in field_mappings if fm.matched_data)
        high_conf_count = sum(1 for fm in field_mappings if fm.confidence and fm.confidence >= 0.8)

        add_matching_log(f"✅ Matching complete!")
        add_matching_log(f"  → Total fields: {len(field_mappings)}")
        add_matching_log(f"  → Matched: {matched_count}")
        add_matching_log(f"  → High confidence (≥0.8): {high_conf_count}")

        with log_display:
            st.code("\n".join(st.session_state.matching_log_messages), language="log")

        # Step 4: Save mappings
        add_matching_log("💾 Saving field mappings to workspace...")
        progress_container.info("💾 Saving field mappings...")

        success = workspace_manager.save_field_mappings(workspace_id, field_mappings)
        if success:
            st.session_state.field_mappings = field_mappings
            add_matching_log(f"✅ Field mappings saved successfully")

            # Update workspace status
            workspace = st.session_state.current_workspace
            workspace.status = WorkspaceStatus.FIELD_MATCHING
            workspace_manager._save_metadata(workspace_id, workspace)

            add_matching_log(f"✅ Workspace status updated")
            add_matching_log("")
            add_matching_log("🎉 Field matching completed successfully!")

            with log_display:
                st.code("\n".join(st.session_state.matching_log_messages), language="log")

            progress_container.success("✅ Matching complete! Review matches below.")
            st.session_state.matching_in_progress = False

            # Small delay to show final logs
            import time
            time.sleep(1)
            st.rerun()
        else:
            add_matching_log("❌ ERROR: Failed to save field mappings")
            with log_display:
                st.code("\n".join(st.session_state.matching_log_messages), language="log")
            progress_container.error("Failed to save field mappings")
            st.session_state.matching_in_progress = False

    except Exception as e:
        add_matching_log(f"❌ ERROR: {str(e)}")
        logger.error(f"Field matching failed: {e}", exc_info=True)

        with log_display:
            st.code("\n".join(st.session_state.matching_log_messages), language="log")

        progress_container.error(f"Field matching failed: {str(e)}")
        st.session_state.matching_in_progress = False


def render_matching_interface():
    """Render the field matching review interface."""
    workspace_id = st.session_state.current_workspace.workspace_id
    workspace_manager = st.session_state.workspace_manager
    field_mappings = st.session_state.field_mappings

    # Progress statistics
    render_matching_statistics(field_mappings)

    st.divider()

    # Filter controls
    col1, col2, col3 = st.columns(3)
    with col1:
        filter_status = st.selectbox(
            "Show:",
            ["All", "Matched", "Unmatched", "Approved", "Pending Review"],
            key="filter_status_select"
        )
        st.session_state.filter_status = filter_status

    with col2:
        filter_confidence = st.slider(
            "Min Confidence:",
            0.0, 1.0, 0.0, 0.05,
            key="filter_confidence_slider"
        )
        st.session_state.filter_confidence = filter_confidence

    with col3:
        st.write("")  # Spacing
        if st.button("🔄 Refresh"):
            st.rerun()

    # Bulk actions
    st.divider()
    render_bulk_actions(field_mappings)

    st.divider()

    # Field list
    render_field_list(field_mappings)

    # Navigation
    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        if st.button("⬅️ Back to Extraction"):
            st.session_state.workflow_step = 3
            st.rerun()
    with col2:
        approved_count = sum(1 for fm in field_mappings if fm.user_approved)
        if approved_count > 0:
            if st.button("Next: Fill & Export ➡️", type="primary"):
                st.session_state.workflow_step = 5
                st.rerun()


def render_matching_statistics(field_mappings):
    """Render progress statistics."""
    total = len(field_mappings)
    matched = sum(1 for fm in field_mappings if fm.matched_data is not None)
    approved = sum(1 for fm in field_mappings if fm.user_approved)
    high_conf = sum(1 for fm in field_mappings if fm.confidence and fm.confidence >= 0.8)

    completion_pct = (approved / total * 100) if total > 0 else 0

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Fields", total)
    with col2:
        st.metric("Matched", f"{matched}/{total}", f"{matched/total*100:.0f}%" if total > 0 else "0%")
    with col3:
        st.metric("Approved", f"{approved}/{total}", f"{completion_pct:.0f}%")
    with col4:
        st.metric("High Confidence", high_conf, f"≥80%")

    # Progress bar
    st.progress(completion_pct / 100, text=f"Approval Progress: {completion_pct:.0f}%")


def render_bulk_actions(field_mappings):
    """Render bulk action buttons."""
    workspace_id = st.session_state.current_workspace.workspace_id
    workspace_manager = st.session_state.workspace_manager

    col1, col2, col3 = st.columns(3)

    with col1:
        high_conf_count = sum(
            1 for fm in field_mappings
            if fm.confidence and fm.confidence >= 0.8 and not fm.user_approved
        )
        if st.button(f"✅ Approve All High Confidence ({high_conf_count})", disabled=high_conf_count == 0):
            for fm in field_mappings:
                if fm.confidence and fm.confidence >= 0.8 and not fm.user_approved:
                    workspace_manager.update_field_mapping(
                        workspace_id,
                        fm.field_id,
                        {"user_approved": True}
                    )
            # Reload mappings
            st.session_state.field_mappings = workspace_manager.get_field_mappings(workspace_id)
            st.success(f"Approved {high_conf_count} high-confidence matches")
            st.rerun()

    with col2:
        matched_count = sum(1 for fm in field_mappings if fm.matched_data is not None and not fm.user_approved)
        if st.button(f"✅ Approve All Matched ({matched_count})", disabled=matched_count == 0):
            for fm in field_mappings:
                if fm.matched_data is not None and not fm.user_approved:
                    workspace_manager.update_field_mapping(
                        workspace_id,
                        fm.field_id,
                        {"user_approved": True}
                    )
            st.session_state.field_mappings = workspace_manager.get_field_mappings(workspace_id)
            st.success(f"Approved {matched_count} matched fields")
            st.rerun()

    with col3:
        approved_count = sum(1 for fm in field_mappings if fm.user_approved)
        if st.button(f"❌ Reset All Approvals ({approved_count})", disabled=approved_count == 0):
            for fm in field_mappings:
                if fm.user_approved:
                    workspace_manager.update_field_mapping(
                        workspace_id,
                        fm.field_id,
                        {"user_approved": False}
                    )
            st.session_state.field_mappings = workspace_manager.get_field_mappings(workspace_id)
            st.success(f"Reset {approved_count} approvals")
            st.rerun()


def render_field_list(field_mappings):
    """Render filtered list of field cards."""
    # Apply filters
    filtered_mappings = apply_filters(field_mappings)

    if not filtered_mappings:
        st.info("No fields match the current filters")
        return

    st.markdown(f"**Showing {len(filtered_mappings)} of {len(field_mappings)} fields**")

    # Render each field card
    for mapping in filtered_mappings:
        render_field_card(mapping)


def apply_filters(field_mappings):
    """Apply current filters to field mappings."""
    filtered = field_mappings

    # Status filter
    if st.session_state.filter_status == "Matched":
        filtered = [fm for fm in filtered if fm.matched_data is not None]
    elif st.session_state.filter_status == "Unmatched":
        filtered = [fm for fm in filtered if fm.matched_data is None]
    elif st.session_state.filter_status == "Approved":
        filtered = [fm for fm in filtered if fm.user_approved]
    elif st.session_state.filter_status == "Pending Review":
        filtered = [fm for fm in filtered if not fm.user_approved]

    # Confidence filter
    if st.session_state.filter_confidence > 0:
        filtered = [
            fm for fm in filtered
            if fm.confidence and fm.confidence >= st.session_state.filter_confidence
        ]

    return filtered


def render_field_card(mapping):
    """Render a single field card with actions."""
    workspace_id = st.session_state.current_workspace.workspace_id
    workspace_manager = st.session_state.workspace_manager

    # Confidence badge color
    if mapping.confidence is None or mapping.confidence == 0:
        badge = "⚪ **UNMATCHED**"
        badge_color = "#999"
    elif mapping.confidence >= 0.95:
        badge = f"🟢 **HIGH** ({mapping.confidence:.2f})"
        badge_color = "#28a745"
    elif mapping.confidence >= 0.7:
        badge = f"🟡 **MEDIUM** ({mapping.confidence:.2f})"
        badge_color = "#ffc107"
    else:
        badge = f"🟠 **LOW** ({mapping.confidence:.2f})"
        badge_color = "#fd7e14"

    # Approval status
    status_emoji = "✅" if mapping.user_approved else "⏳"

    with st.expander(f"{status_emoji} **{mapping.field_description}** - {badge}", expanded=not mapping.user_approved):
        col1, col2 = st.columns([2, 1])

        with col1:
            st.markdown(f"**Location:** {mapping.field_location}")
            if mapping.field_type:
                st.markdown(f"**Type:** `{mapping.field_type}`")

            if mapping.matched_data:
                st.markdown("**Matched Data:**")
                # Check if user is editing
                if f"editing_{mapping.field_id}" in st.session_state and st.session_state[f"editing_{mapping.field_id}"]:
                    new_value = st.text_area(
                        "Edit value:",
                        value=mapping.user_override or mapping.matched_data,
                        key=f"edit_text_{mapping.field_id}",
                        height=100
                    )
                    col_a, col_b = st.columns(2)
                    with col_a:
                        if st.button("💾 Save", key=f"save_{mapping.field_id}"):
                            workspace_manager.update_field_mapping(
                                workspace_id,
                                mapping.field_id,
                                {"user_override": new_value, "user_approved": True}
                            )
                            st.session_state[f"editing_{mapping.field_id}"] = False
                            st.session_state.field_mappings = workspace_manager.get_field_mappings(workspace_id)
                            st.success("Saved!")
                            st.rerun()
                    with col_b:
                        if st.button("❌ Cancel", key=f"cancel_{mapping.field_id}"):
                            st.session_state[f"editing_{mapping.field_id}"] = False
                            st.rerun()
                else:
                    display_value = mapping.user_override or mapping.matched_data
                    st.markdown(f"```\n{display_value}\n```")
                    if mapping.user_override:
                        st.caption("(User edited)")

                if mapping.data_source:
                    st.markdown(f"**Source:** `{mapping.data_source}`")

            else:
                st.warning("No match found")

        with col2:
            if mapping.matched_data and not mapping.user_approved:
                if st.button("✅ Approve", key=f"approve_{mapping.field_id}", use_container_width=True):
                    workspace_manager.update_field_mapping(
                        workspace_id,
                        mapping.field_id,
                        {"user_approved": True}
                    )
                    st.session_state.field_mappings = workspace_manager.get_field_mappings(workspace_id)
                    st.rerun()

            if mapping.matched_data:
                if f"editing_{mapping.field_id}" not in st.session_state:
                    st.session_state[f"editing_{mapping.field_id}"] = False

                if not st.session_state[f"editing_{mapping.field_id}"]:
                    if st.button("✏️ Edit", key=f"edit_{mapping.field_id}", use_container_width=True):
                        st.session_state[f"editing_{mapping.field_id}"] = True
                        st.rerun()

            if mapping.user_approved:
                if st.button("↩️ Unapprove", key=f"unapprove_{mapping.field_id}", use_container_width=True):
                    workspace_manager.update_field_mapping(
                        workspace_id,
                        mapping.field_id,
                        {"user_approved": False}
                    )
                    st.session_state.field_mappings = workspace_manager.get_field_mappings(workspace_id)
                    st.rerun()


def render_step5_fill_export():
    """Step 5: Fill tender and export."""
    section_header("✅", "Step 5: Fill & Export", "Complete tender document")

    st.info("🚧 **Phase 3 - Coming Soon**")
    st.markdown("""
    This step will:
    1. Fill tender document with matched data
    2. Maintain original formatting
    3. Allow manual edits
    4. Export completed tender document

    **Status:** Phase 3 in planning 📋
    """)


def main():
    """Main page function."""
    st.title("📝 Proposal Workflow")
    st.caption(f"Version: {PAGE_VERSION}")

    st.markdown("""
    Complete workflow for responding to tenders with workspace-based extraction.
    Each tender gets its own workspace for progressive document building.
    """)

    # Initialize session state
    initialize_session_state()

    # Initialize workspace manager
    config_manager = ConfigManager()
    config = config_manager.get_config()
    db_path = config.get('ai_database_path')

    if not db_path:
        st.error("❌ Database path not configured. Please set it in the main settings.")
        return

    # Convert path
    wsl_db_path = convert_windows_to_wsl_path(db_path)

    try:
        if st.session_state.workspace_manager is None:
            st.session_state.workspace_manager = WorkspaceManager(Path(wsl_db_path))
            logger.info("Workspace manager initialized")
    except Exception as e:
        st.error(f"❌ Failed to initialize workspace manager: {str(e)}")
        logger.error(f"Workspace manager initialization failed: {e}", exc_info=True)
        return

    # Show current workspace info if exists
    if st.session_state.current_workspace:
        workspace = st.session_state.current_workspace
        with st.expander("📁 Current Workspace", expanded=False):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.caption(f"**ID:** {workspace.workspace_id}")
            with col2:
                st.caption(f"**Tender:** {workspace.tender_id}")
            with col3:
                st.caption(f"**Status:** {workspace.status.replace('_', ' ').title()}")

    # Show progress
    st.divider()
    render_workflow_progress(st.session_state.workflow_step)
    st.divider()

    # Render current step
    if st.session_state.workflow_step == 1:
        render_step1_select_tender()
    elif st.session_state.workflow_step == 2:
        render_step2_select_sources()
    elif st.session_state.workflow_step == 3:
        render_step3_extract_data()
    elif st.session_state.workflow_step == 4:
        render_step4_match_fields()
    elif st.session_state.workflow_step == 5:
        render_step5_fill_export()

    # Reset button
    st.divider()
    col1, col2 = st.columns([3, 1])

    with col1:
        if st.button("🔄 Start Over"):
            st.session_state.workflow_step = 1
            st.session_state.tender_document = None
            st.session_state.tender_filename = None
            st.session_state.tender_id = None
            st.session_state.source_documents = []
            st.session_state.extracted_data = None
            st.session_state.selected_entity = None
            st.session_state.current_workspace = None
            st.rerun()

    with col2:
        if st.session_state.current_workspace and st.button("🗑️ Delete Workspace"):
            workspace_id = st.session_state.current_workspace.workspace_id
            if st.session_state.workspace_manager.delete_workspace(workspace_id):
                st.success(f"Deleted workspace: {workspace_id}")
                st.session_state.current_workspace = None
                st.rerun()


if __name__ == "__main__":
    main()
