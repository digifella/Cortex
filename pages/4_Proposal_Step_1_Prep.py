# ## File: pages/5_Proposal_Step_1.py
# Version: v4.7.0
# Date: 2025-07-15
# Purpose: An interactive UI to tag document placeholders.
#          - CRITICAL FIX (v27.0.0): The parsing logic has been entirely rewritten
#            to correctly handle placeholder paragraphs. The editor now identifies
#            a heading, and then presents EVERY subsequent paragraph (including empty
#            ones) as a taggable item, resolving the "missing sections" bug.
#          - FIX (v27.0.0): The template processing logic was updated to match the
#            new parsing logic, ensuring correct replacement of all items.

import streamlit as st
import sys
from pathlib import Path
import io
import docx
from docx.text.paragraph import Paragraph
from docx.table import Table
import json
import os

# --- Project Setup ---
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from cortex_engine.instruction_parser import iter_block_items

# --- Configuration ---
st.set_page_config(layout="wide", page_title="Cortex Template Editor")

MAPS_DIR = project_root / "template_maps"
MAPS_DIR.mkdir(exist_ok=True)

CORTEX_ACTIONS = {
    "Ignore this block": "[IGNORE]",
    "Insert Content from a File": "[INSERT_FROM_FILE]",
    "Human Input (Manual Text)": "[PROMPT_HUMAN]",
    "AI: Generate from KB only": "[GENERATE_FROM_KB_ONLY]",
    "AI: Generate from Proposal only": "[GENERATE_FROM_PROPOSAL_ONLY]",
    "AI: Generate from both KB and Proposal": "[GENERATE_FROM_KB_AND_PROPOSAL]",
    "AI: Retrieve Case Studies from KB": "[RETRIEVE_FROM_KB]",
    "AI: Generate Resources": "[GENERATE_RESOURCES]",
}


# --- Core Functions ---

def initialize_state():
    """Initializes all session state keys needed for the editor."""
    keys = ['template_doc_bytes', 'processed_doc_bytes', 'template_elements', 'action_map', 'template_filename', 'inserted_content']
    for key in keys:
        if key not in st.session_state:
            st.session_state[key] = None if 'bytes' in key else [] if 'elements' in key else {} if 'map' in key or 'content' in key else ""

def parse_document(doc_bytes):
    """
    (RE-ARCHITECTED) Iterates through all document blocks (paragraphs and tables).
    It identifies headings to establish the current section context, and then
    presents all subsequent paragraphs (including empty ones) and table cells as taggable items.
    """
    doc = docx.Document(io.BytesIO(doc_bytes))
    elements_for_ui = []
    action_map = {}
    current_heading = "Document Header (No Section Found Yet)"

    for i, block in enumerate(iter_block_items(doc)):
        block_id = f"block_{i}"

        if isinstance(block, Paragraph):
            p = block
            text = p.text.strip()
            # A heading is defined by style or by a specific text convention.
            is_heading = p.style.name.startswith('Heading') or text.lstrip(' "*•').startswith('Section:')

            if is_heading:
                current_heading = text if text else "Unnamed Heading"
                continue  # Skip adding headings to the UI, they just provide context.

            # Any paragraph that IS NOT a heading is a taggable placeholder.
            # This includes empty paragraphs which are crucial for defining insertion points.
            ui_content = text if text else "[Empty paragraph available for tagging]"
            elements_for_ui.append({
                "id": block_id,
                "content": ui_content,
                "heading": current_heading,
                "type": "Paragraph"
            })
            action_map[block_id] = "[IGNORE]"

        elif isinstance(block, Table):
            for r_idx, row in enumerate(block.rows):
                for c_idx, cell in enumerate(row.cells):
                    cell_id = f"{block_id}_r{r_idx}_c{c_idx}"
                    content = cell.text.strip()
                    # Show any cell that has content or is an empty response cell in a Q&A table.
                    if content or (c_idx > 0 and row.cells[c_idx-1].text.strip()):
                        ui_content = content if content else f"(Response for: '{row.cells[c_idx-1].text.strip()[:40]}...')"
                        elements_for_ui.append({
                            "id": cell_id,
                            "content": ui_content,
                            "heading": f"{current_heading} - Table Row {r_idx+1}",
                            "type": "Table Cell"
                        })
                        action_map[cell_id] = "[IGNORE]"

    return elements_for_ui, action_map

def process_template():
    """
    (CORRECTED) Assembles the final template by iterating through the document
    and replacing content based on the action map. This logic mirrors `parse_document`.
    """
    if not st.session_state.template_doc_bytes:
        st.error("No template loaded.")
        return

    doc = docx.Document(io.BytesIO(st.session_state.template_doc_bytes))

    for i, block in enumerate(iter_block_items(doc)):
        block_id = f"block_{i}"

        if isinstance(block, Paragraph):
            p = block
            text = p.text.strip()
            is_heading = p.style.name.startswith('Heading') or text.lstrip(' "*•').startswith('Section:')

            if is_heading:
                continue # Never modify headings

            # This paragraph is a taggable placeholder. Check if an action applies.
            action = st.session_state.action_map.get(block_id)
            if action and action != "[IGNORE]":
                p.clear()  # Clear all runs (content and formatting)
                if action == "[INSERT_FROM_FILE]":
                    content_to_insert = st.session_state.inserted_content.get(block_id, f"[{action} - ERROR: NO FILE UPLOADED]")
                    p.add_run(content_to_insert)
                else:
                    p.add_run(action)

        elif isinstance(block, Table):
            for r_idx, row in enumerate(block.rows):
                for c_idx, cell in enumerate(row.cells):
                    cell_id = f"{block_id}_r{r_idx}_c{c_idx}"
                    action = st.session_state.action_map.get(cell_id)
                    if action and action != "[IGNORE]":
                        cell.text = ""  # Clear cell content
                        if action == "[INSERT_FROM_FILE]":
                            content_to_insert = st.session_state.inserted_content.get(cell_id, f"[{action} - ERROR: NO FILE UPLOADED]")
                            cell.text = content_to_insert
                        else:
                            cell.text = action
    bio = io.BytesIO()
    doc.save(bio)
    st.session_state.processed_doc_bytes = bio.getvalue()
    st.success("✅ Template processed successfully!")


# --- UI Callbacks ---

def handle_upload():
    uploaded_file = st.session_state.get('template_uploader')
    if uploaded_file:
        initialize_state()
        st.session_state.template_doc_bytes = uploaded_file.getvalue()
        st.session_state.template_filename = uploaded_file.name
        with st.spinner("Parsing document to identify all taggable content..."):
            elements, actions = parse_document(st.session_state.template_doc_bytes)
        st.session_state.template_elements = elements
        st.session_state.action_map = actions

def handle_file_insertion(element_id):
    """Reads content from uploaded file and stores it in session state."""
    uploader_key = f"uploader_{element_id}"
    uploaded_file = st.session_state.get(uploader_key)
    if uploaded_file:
        try:
            if uploaded_file.name.endswith('.txt'):
                content = uploaded_file.getvalue().decode('utf-8')
            else:  # .docx
                doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
                content = "\n".join([p.text for p in doc.paragraphs])
            st.session_state.inserted_content[element_id] = content
        except Exception as e:
            st.error(f"Failed to read file: {e}")

def save_progress():
    if not st.session_state.template_filename:
        st.error("Cannot save without an uploaded file.")
        return
    save_data = {
        "source_template_filename": st.session_state.template_filename,
        "action_map": st.session_state.action_map,
        "inserted_content": st.session_state.inserted_content,
    }
    save_path = MAPS_DIR / f"{Path(st.session_state.template_filename).stem}_mapping.json"
    try:
        with open(save_path, 'w', encoding='utf-8') as f: json.dump(save_data, f, indent=4)
        st.success(f"✅ Progress saved to `{save_path}`")
    except Exception as e: st.error(f"Failed to save progress: {e}")

def load_progress(map_file_path):
    try:
        with open(map_file_path, 'r', encoding='utf-8') as f: load_data = json.load(f)
        if load_data.get("source_template_filename") != st.session_state.template_filename:
            st.warning("Warning: This mapping file appears to be for a different source document.")
        st.session_state.action_map = load_data.get("action_map", {})
        st.session_state.inserted_content = load_data.get("inserted_content", {})
        st.success(f"✅ Progress loaded from `{Path(map_file_path).name}`!")
    except Exception as e: st.error(f"Failed to load progress file: {e}")

# --- Streamlit UI Layout ---
st.title("📝 6. Template Editor")
st.caption("Version 27.0.0 - Corrected Placeholder Parsing")

initialize_state()

with st.expander("Start Here: Upload Your Template", expanded=True):
    st.info("Upload any `.docx` file. The editor will show all content, allowing you to replace any part with a Cortex instruction or insert boilerplate from a file.")
    st.file_uploader("Upload document", type="docx", key="template_uploader", on_change=handle_upload)
    st.header("Load Saved Progress")
    saved_maps = [f for f in os.listdir(MAPS_DIR) if f.endswith('_mapping.json')]
    if saved_maps:
        selected_map_file = st.selectbox("Select progress file", options=[""] + saved_maps, index=0, help="To load progress, you must first upload the matching template file.")
        if selected_map_file and st.session_state.template_doc_bytes:
            if st.button("Load Selected Progress"): load_progress(MAPS_DIR / selected_map_file)

if st.session_state.template_elements:
    st.divider()
    st.header("Assign Cortex Actions to Document Content")
    st.markdown("For any content you want to keep as-is, leave the action as **'Ignore this block'**. To insert an instruction, select an action from the dropdown to **replace** the original text.")
    st.button("💾 Save Progress", on_click=save_progress)

    for element in st.session_state.template_elements:
        elem_id = element['id']
        heading = element['heading']
        content = element['content']

        st.markdown("---")
        st.markdown(f"**Found {element['type']} under Section:** `{heading}`")
        st.info(f"**Original Text:**\n\n```\n{content}\n```")

        action_keys = list(CORTEX_ACTIONS.keys())
        current_action_tag = st.session_state.action_map.get(elem_id, "[IGNORE]")

        try:
            current_text_key = next(k for k, v in CORTEX_ACTIONS.items() if v == current_action_tag)
            default_index = action_keys.index(current_text_key)
        except (StopIteration, ValueError): default_index = 0

        selected_text = st.selectbox("Action for this block:", options=action_keys, index=default_index, key=f"action_{elem_id}")
        st.session_state.action_map[elem_id] = CORTEX_ACTIONS[selected_text]

        if st.session_state.action_map[elem_id] == "[INSERT_FROM_FILE]":
            st.file_uploader(
                "Upload boilerplate (.txt or .docx)",
                type=['txt', 'docx'],
                key=f"uploader_{elem_id}",
                on_change=handle_file_insertion,
                args=(elem_id,)
            )
            if elem_id in st.session_state.inserted_content:
                st.success("Boilerplate file loaded and ready for insertion.")

    st.divider()
    st.header("Generate Final Tagged Template")
    st.info("This creates a new `.docx` file where any content you tagged is replaced with the chosen Cortex instruction or boilerplate text.")
    st.button("⚙️ Process and Generate Tagged Template", on_click=process_template, type="primary", use_container_width=True)

    if st.session_state.processed_doc_bytes:
        new_filename = f"CORTEX_TPL_{st.session_state.template_filename}"
        st.download_button(label=f"⬇️ Download '{new_filename}'", data=st.session_state.processed_doc_bytes, file_name=new_filename)
else:
    st.markdown("---")
    st.warning("Upload a document to begin.")